"""
Meeting minutes export API.
"""
from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any, Optional
import io
import os
import re
import tempfile
import unicodedata

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from fastapi import APIRouter, Depends, HTTPException, Response
from fpdf import FPDF
from fpdf.errors import FPDFException
from pydantic import BaseModel
from sqlalchemy.orm import Session

from . import auth, models
from .database import get_db

router = APIRouter(prefix="/api/export", tags=["export"])
EXPORT_FILENAME_PATTERN = re.compile(
    r"^([a-z0-9-]+)_([a-z0-9-]+)_(\d{2}-\d{2}-\d{4}_\d{2}-\d{2})\.(pdf|docx)$"
)

SUPPORTED_LANGUAGES = {"vi", "en", "ja", "zh", "ko"}
PLACEHOLDER_TITLES = {
    "cuoc hop",
    "cuoc-hop",
    "meeting",
    "untitled",
    "new meeting",
    "new-meeting",
}
SERIF_FONT_CANDIDATES = [
    (
        "/System/Library/Fonts/Supplemental/Times New Roman.ttf",
        "/System/Library/Fonts/Supplemental/Times New Roman Bold.ttf",
    ),
    (
        "/System/Library/Fonts/Times.ttc",
        "/System/Library/Fonts/Times.ttc",
    ),
]
UNICODE_FONT_CANDIDATES = [
    (
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
    ),
    (
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
    ),
    (
        "/Library/Fonts/Arial Unicode.ttf",
        "/Library/Fonts/Arial Unicode.ttf",
    ),
]

LANGUAGE_LABELS: dict[str, dict[str, str]] = {
    "vi": {
        "document_title": "BIÊN BẢN CUỘC HỌP",
        "document_subtitle": "Mẫu biên bản doanh nghiệp",
        "meeting_name": "Tên cuộc họp",
        "meeting_time": "Thời gian",
        "duration": "Thời lượng",
        "location_mode": "Địa điểm / hình thức",
        "participants": "Thành phần tham dự",
        "main_content": "Nội dung chính",
        "summary": "Tóm tắt cuộc họp",
        "key_points": "Nội dung nổi bật",
        "decisions": "Quyết định / thống nhất",
        "action_items": "Công việc cần thực hiện",
        "conclusion": "Kết luận",
        "appendix_title": "PHỤ LỤC A. TRANSCRIPT CUỘC HỌP",
        "appendix_note": "Transcript được đính kèm để đối chiếu nội dung chi tiết theo từng mốc thời gian.",
        "chairperson": "Người chủ trì",
        "secretary": "Thư ký",
        "confirmation": "Xác nhận / Người tham dự",
        "signature_placeholder": "(Ký, ghi rõ họ tên)",
        "no_data": "Chưa có dữ liệu",
        "owner": "Phụ trách",
        "deadline": "Hạn",
        "status": "Trạng thái",
        "priority": "Ưu tiên",
        "date_format": "%H giờ %M phút, ngày %d tháng %m năm %Y",
        "duration_unit": "phút",
        "meeting_mode_fallback": "Không xác định",
        "generated_at": "Thời điểm xuất",
        "signature_section": "Xác nhận",
        "unknown_title": "Chưa cập nhật tên cuộc họp",
        "time_start": "Bắt đầu",
        "time_end": "Kết thúc",
        "meeting_creator": "Người tạo cuộc họp",
        "time_partial_separator": " | ",
        "organization_name": "Đơn vị tổ chức",
        "meeting_description": "Mô tả",
        "meeting_type_label": "Loại cuộc họp",
        "action_table_no": "STT",
        "action_table_title": "Nội dung",
        "action_table_owner": "Phụ trách",
        "action_table_deadline": "Hạn",
        "action_table_status": "Trạng thái",
        "action_table_priority": "Ưu tiên",
        "risks": "Rủi ro / Điểm nghẽn",
        "open_questions": "Câu hỏi mở",
        "timeline_highlights": "Dòng thời gian nổi bật",
        "speaker_summaries": "Tóm tắt theo người nói",
        "download_date": "Ngày xuất biên bản",
    },
    "en": {
        "document_title": "MEETING MINUTES",
        "document_subtitle": "Corporate meeting minutes",
        "meeting_name": "Meeting name",
        "meeting_time": "Time",
        "duration": "Duration",
        "location_mode": "Location / format",
        "participants": "Participants",
        "main_content": "Main content",
        "summary": "Meeting summary",
        "key_points": "Key points",
        "decisions": "Decisions",
        "action_items": "Action items",
        "conclusion": "Conclusion",
        "appendix_title": "APPENDIX A. MEETING TRANSCRIPT",
        "appendix_note": "Transcript appendix for detailed line-by-line review.",
        "chairperson": "Chairperson",
        "secretary": "Secretary",
        "confirmation": "Attendee confirmation",
        "signature_placeholder": "(Signature and full name)",
        "no_data": "No data available",
        "owner": "Owner",
        "deadline": "Deadline",
        "status": "Status",
        "priority": "Priority",
        "date_format": "%Y-%m-%d %H:%M",
        "duration_unit": "minutes",
        "meeting_mode_fallback": "Not specified",
        "generated_at": "Generated at",
        "signature_section": "Sign-off",
        "unknown_title": "Meeting title not updated",
        "time_start": "Start",
        "time_end": "End",
        "meeting_creator": "Meeting creator",
        "time_partial_separator": " | ",
        "organization_name": "Organization",
        "meeting_description": "Description",
        "meeting_type_label": "Meeting type",
        "action_table_no": "No.",
        "action_table_title": "Title",
        "action_table_owner": "Owner",
        "action_table_deadline": "Deadline",
        "action_table_status": "Status",
        "action_table_priority": "Priority",
        "risks": "Risks / Blockers",
        "open_questions": "Open Questions",
        "timeline_highlights": "Timeline Highlights",
        "speaker_summaries": "Speaker Summaries",
        "download_date": "Export Date",
    },
    "ja": {
        "document_title": "会議議事録",
        "document_subtitle": "企業向け会議議事録",
        "meeting_name": "会議名",
        "meeting_time": "日時",
        "duration": "所要時間",
        "location_mode": "場所 / 形式",
        "participants": "出席者",
        "main_content": "主な内容",
        "summary": "会議要約",
        "key_points": "重要ポイント",
        "decisions": "決定事項",
        "action_items": "対応事項",
        "conclusion": "結論",
        "appendix_title": "付録A. 会議トランスクリプト",
        "appendix_note": "詳細確認用のトランスクリプト付録です。",
        "chairperson": "議長",
        "secretary": "書記",
        "confirmation": "出席確認",
        "signature_placeholder": "（署名・氏名）",
        "no_data": "データがありません",
        "owner": "担当",
        "deadline": "期限",
        "status": "状態",
        "priority": "優先度",
        "date_format": "%Y/%m/%d %H:%M",
        "duration_unit": "分",
        "meeting_mode_fallback": "未設定",
        "generated_at": "出力日時",
        "signature_section": "確認",
        "unknown_title": "会議名未設定",
        "time_start": "開始",
        "time_end": "終了",
        "meeting_creator": "会議作成者",
        "time_partial_separator": " | ",
        "organization_name": "主催組織",
        "meeting_description": "説明",
        "meeting_type_label": "会議種別",
        "action_table_no": "番号",
        "action_table_title": "内容",
        "action_table_owner": "担当",
        "action_table_deadline": "期限",
        "action_table_status": "状態",
        "action_table_priority": "優先度",
        "risks": "リスク / ボトルネック",
        "open_questions": "未解決の質問",
        "timeline_highlights": "タイムラインハイライト",
        "speaker_summaries": "発言者別要約",
        "download_date": "出力日",
    },
    "zh": {
        "document_title": "会议纪要",
        "document_subtitle": "企业会议纪要",
        "meeting_name": "会议名称",
        "meeting_time": "时间",
        "duration": "时长",
        "location_mode": "地点 / 形式",
        "participants": "参会人员",
        "main_content": "主要内容",
        "summary": "会议总结",
        "key_points": "要点",
        "decisions": "决定事项",
        "action_items": "待办事项",
        "conclusion": "结论",
        "appendix_title": "附录A. 会议转写",
        "appendix_note": "用于详细核对的 transcript 附录。",
        "chairperson": "主持人",
        "secretary": "秘书",
        "confirmation": "参会确认",
        "signature_placeholder": "（签名及姓名）",
        "no_data": "暂无数据",
        "owner": "负责人",
        "deadline": "截止日",
        "status": "状态",
        "priority": "优先级",
        "date_format": "%Y-%m-%d %H:%M",
        "duration_unit": "分钟",
        "meeting_mode_fallback": "未指定",
        "generated_at": "导出时间",
        "signature_section": "确认",
        "unknown_title": "会议名称未更新",
        "time_start": "开始",
        "time_end": "结束",
        "meeting_creator": "会议创建人",
        "time_partial_separator": " | ",
        "organization_name": "主办单位",
        "meeting_description": "描述",
        "meeting_type_label": "会议类型",
        "action_table_no": "序号",
        "action_table_title": "内容",
        "action_table_owner": "负责人",
        "action_table_deadline": "截止日",
        "action_table_status": "状态",
        "action_table_priority": "优先级",
        "risks": "风险 / 瓶颈",
        "open_questions": "未解决的问题",
        "timeline_highlights": "时间线亮点",
        "speaker_summaries": "发言人摘要",
        "download_date": "导出日期",
    },
    "ko": {
        "document_title": "회의록",
        "document_subtitle": "기업 회의 문서",
        "meeting_name": "회의명",
        "meeting_time": "일시",
        "duration": "소요 시간",
        "location_mode": "장소 / 형식",
        "participants": "참석자",
        "main_content": "주요 내용",
        "summary": "회의 요약",
        "key_points": "핵심 포인트",
        "decisions": "결정 사항",
        "action_items": "실행 과제",
        "conclusion": "결론",
        "appendix_title": "부록 A. 회의 트랜스크립트",
        "appendix_note": "상세 확인용 transcript 부록입니다.",
        "chairperson": "주재",
        "secretary": "간서",
        "confirmation": "참석 확인",
        "signature_placeholder": "(서명 및 성명)",
        "no_data": "데이터 없음",
        "owner": "담당자",
        "deadline": "마감",
        "status": "상태",
        "priority": "우선순위",
        "date_format": "%Y-%m-%d %H:%M",
        "duration_unit": "분",
        "meeting_mode_fallback": "미지정",
        "generated_at": "내보낸 시각",
        "signature_section": "확인",
        "unknown_title": "회의명이 아직 업데이트되지 않음",
        "time_start": "시작",
        "time_end": "종료",
        "meeting_creator": "회의 생성자",
        "time_partial_separator": " | ",
        "organization_name": "주최 기관",
        "meeting_description": "설명",
        "meeting_type_label": "회의 유형",
        "action_table_no": "번호",
        "action_table_title": "내용",
        "action_table_owner": "담당자",
        "action_table_deadline": "마감",
        "action_table_status": "상태",
        "action_table_priority": "우선순위",
        "risks": "위험 / 병목",
        "open_questions": "미해결 질문",
        "timeline_highlights": "타임라인 하이라이트",
        "speaker_summaries": "발언자별 요약",
        "download_date": "내보내기 날짜",
    },
}


class ExportRequest(BaseModel):
    meeting_id: str
    format: str  # "pdf" or "docx"
    language: str = "vi"
    include_transcript_appendix: Optional[bool] = None
    include_transcript: bool = True
    include_summary: bool = True
    include_action_items: bool = True


class ExportResponse(BaseModel):
    download_url: str
    filename: str
    size_bytes: int
    created_at: datetime


def _format_action_item_assignee_summary(item: models.ActionItem) -> str:
    assignees = getattr(item, "assignees", []) or []
    if assignees:
        return ", ".join(
            (assignee.display_name or assignee.email or "N/A")
            for assignee in assignees
        )
    return item.assigned_email or getattr(item.assigned_to_user, "display_name", None) or item.assigned_to or "N/A"


def _normalize_language(language: Optional[str]) -> str:
    normalized = (language or "vi").lower()
    return normalized if normalized in SUPPORTED_LANGUAGES else "vi"


def _labels(language: str) -> dict[str, str]:
    return LANGUAGE_LABELS[_normalize_language(language)]


def _wants_transcript_appendix(request: ExportRequest) -> bool:
    if request.include_transcript_appendix is not None:
        return request.include_transcript_appendix
    return request.include_transcript


def _slugify_filename_part(value: str) -> str:
    ascii_text = unicodedata.normalize("NFKD", value or "").encode("ascii", "ignore").decode("ascii")
    lowered = ascii_text.strip().lower()
    normalized = re.sub(r"[^a-z0-9]+", "-", lowered)
    collapsed = re.sub(r"-{2,}", "-", normalized).strip("-")
    return collapsed or "cuoc-hop"


def _normalize_title_token(value: Optional[str]) -> str:
    return _slugify_filename_part(value or "")


def _is_placeholder_title(value: Optional[str]) -> bool:
    token = _normalize_title_token(value)
    return not token or token in PLACEHOLDER_TITLES


def _extract_list(value: Any) -> list[str]:
    if not value:
        return []
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    if isinstance(value, tuple):
        return [str(item).strip() for item in value if str(item).strip()]
    if isinstance(value, dict):
        results: list[str] = []
        for item in value.values():
            if isinstance(item, list):
                results.extend(_extract_list(item))
            elif str(item).strip():
                results.append(str(item).strip())
        return results
    text = str(value).strip()
    return [text] if text else []


def _pick_summary(meeting_id: str, language: str, db: Session) -> Optional[models.MeetingSummary]:
    summaries = (
        db.query(models.MeetingSummary)
        .filter(models.MeetingSummary.meeting_id == meeting_id)
        .order_by(models.MeetingSummary.created_at.desc())
        .all()
    )
    if not summaries:
        return None
    preferred = [
        summary
        for summary in summaries
        if summary.language == language and summary.processing_status == "COMPLETED"
    ]
    if preferred:
        return preferred[0]
    preferred = [summary for summary in summaries if summary.language == language]
    if preferred:
        return preferred[0]
    completed = [summary for summary in summaries if summary.processing_status == "COMPLETED"]
    return completed[0] if completed else summaries[0]


def _pick_transcript(meeting_id: str, language: str, db: Session) -> Optional[models.Transcript]:
    transcripts = (
        db.query(models.Transcript)
        .filter(models.Transcript.meeting_id == meeting_id)
        .order_by(models.Transcript.created_at.desc())
        .all()
    )
    if not transcripts:
        return None
    preferred = [
        transcript
        for transcript in transcripts
        if transcript.language == language and transcript.processing_status == "COMPLETED"
    ]
    if preferred:
        return preferred[0]
    preferred = [transcript for transcript in transcripts if transcript.language == language]
    if preferred:
        return preferred[0]
    completed = [transcript for transcript in transcripts if transcript.processing_status == "COMPLETED"]
    return completed[0] if completed else transcripts[0]


def _format_meeting_datetime(value: Optional[datetime], labels: dict[str, str]) -> str:
    if not value:
        return labels["no_data"]
    return value.strftime(labels["date_format"])


def _format_duration(duration_minutes: Optional[int], labels: dict[str, str]) -> str:
    if not duration_minutes:
        return f"0 {labels['duration_unit']}"
    return f"{duration_minutes} {labels['duration_unit']}"


def _format_location_or_mode(meeting: models.Meeting, labels: dict[str, str]) -> str:
    if meeting.location and meeting.location.strip():
        return meeting.location.strip()
    # Fall back to meeting type label
    type_label = _meeting_type_label(meeting, labels)
    if type_label and type_label != labels.get("meeting_mode_fallback", ""):
        return type_label
    # Fall back to organization name
    org_name = _clean_value(getattr(meeting.organization, "name", None)) if meeting.organization else ""
    if org_name:
        return org_name
    return labels["meeting_mode_fallback"]


def _clean_value(value: Optional[str]) -> str:
    return (value or "").strip()


def _participant_display_name(participant: models.MeetingParticipant) -> str:
    user = participant.user
    full_name = " ".join(
        part for part in [getattr(user, "first_name", None), getattr(user, "last_name", None)] if part
    ).strip()
    return (
        _clean_value(participant.name)
        or full_name
        or _clean_value(getattr(user, "username", None))
        or _clean_value(participant.email)
        or _clean_value(getattr(user, "email", None))
        or "Thành viên"
    )


def _participant_business_label(participant: models.MeetingParticipant, labels: dict[str, str]) -> str:
    name = _participant_display_name(participant)
    email = _clean_value(participant.email or getattr(participant.user, "email", None))
    role = _clean_value(participant.role)
    role_suffix = ""
    if role and role.upper() not in ("PARTICIPANT", ""):
        role_map = {
            "HOST": "Chủ trì" if labels.get("meeting_mode_fallback") != "Not specified" else "Host",
            "PRESENTER": "Trình bày" if labels.get("meeting_mode_fallback") != "Not specified" else "Presenter",
            "MODERATOR": "Điều phối" if labels.get("meeting_mode_fallback") != "Not specified" else "Moderator",
            "SECRETARY": labels.get("secretary", "Secretary"),
        }
        role_suffix = f" ({role_map.get(role.upper(), role)})"
    if email and email.lower() != name.lower():
        return f"{name}{role_suffix} - {email}"
    return f"{name}{role_suffix}"


def _creator_label(meeting: models.Meeting, labels: dict[str, str]) -> Optional[str]:
    creator = meeting.created_by_user
    if not creator:
        return None
    name = _clean_value(getattr(creator, "display_name", None)) or " ".join(
        part for part in [getattr(creator, "first_name", None), getattr(creator, "last_name", None)] if part
    ).strip() or _clean_value(getattr(creator, "username", None)) or _clean_value(getattr(creator, "email", None))
    email = _clean_value(getattr(creator, "email", None))
    if not name and not email:
        return None
    if email and email.lower() != name.lower():
        return f"{name} - {email} ({labels['meeting_creator']})"
    return f"{name or email} ({labels['meeting_creator']})"


def _extract_participants(meeting: models.Meeting, labels: dict[str, str]) -> list[str]:
    seen: set[str] = set()
    results: list[str] = []
    participants = sorted(
        meeting.participants or [],
        key=lambda item: (_participant_display_name(item) or "").lower(),
    )
    for participant in participants:
        label = _participant_business_label(participant, labels)
        key = _normalize_title_token(label)
        if label and key and key not in seen:
            seen.add(key)
            results.append(label)

    if not results:
        creator = _creator_label(meeting, labels)
        if creator:
            results.append(creator)

    return results or [labels["no_data"]]


def _display_meeting_title(meeting: models.Meeting, labels: dict[str, str]) -> str:
    title = _clean_value(meeting.title)
    if title:
        return title
    return labels["unknown_title"]


def _filename_title_slug(meeting: models.Meeting) -> str:
    title = _clean_value(meeting.title)
    if title:
        return _slugify_filename_part(title)
    return "cuoc-hop"


def _display_meeting_time(meeting: models.Meeting, labels: dict[str, str]) -> str:
    start = meeting.actual_start or meeting.scheduled_start or meeting.created_at
    end = meeting.actual_end or meeting.scheduled_end
    parts: list[str] = []
    if start:
        parts.append(f"{labels['time_start']}: {_format_meeting_datetime(start, labels)}")
    else:
        parts.append(f"{labels['time_start']}: {labels['no_data']}")
    if end:
        parts.append(f"{labels['time_end']}: {_format_meeting_datetime(end, labels)}")
    elif meeting.actual_start or meeting.scheduled_start:
        parts.append(f"{labels['time_end']}: {labels['no_data']}")
    if meeting.duration:
        parts.append(f"{labels['duration']}: {_format_duration(meeting.duration, labels)}")
    return labels["time_partial_separator"].join(parts)


def _meeting_time_parts(meeting: models.Meeting, labels: dict[str, str]) -> dict[str, str]:
    start = meeting.actual_start or meeting.scheduled_start or meeting.created_at
    end = meeting.actual_end or meeting.scheduled_end
    return {
        "start": _format_meeting_datetime(start, labels) if start else labels["no_data"],
        "end": _format_meeting_datetime(end, labels) if end else labels["no_data"],
        "duration": _format_duration(meeting.duration, labels),
    }


def _meeting_type_label(meeting: models.Meeting, labels: dict[str, str]) -> str:
    is_vi = labels.get("meeting_mode_fallback") != "Not specified"
    type_map = {
        "MEETING": "Cuộc họp" if is_vi else "Meeting",
        "INTERVIEW": "Phỏng vấn" if is_vi else "Interview",
        "TRAINING": "Đào tạo" if is_vi else "Training",
        "REVIEW": "Đánh giá" if is_vi else "Review",
    }
    return type_map.get(meeting.meeting_type or "", labels.get("meeting_mode_fallback", ""))


def _build_speaker_lookup(meeting: models.Meeting) -> dict[str, str]:
    mapping: dict[str, str] = {}
    for speaker in meeting.speaker_mappings or []:
        if speaker.speaker_label and speaker.display_name:
            mapping[speaker.speaker_label] = speaker.display_name
    for participant in meeting.participants or []:
        if participant.speaker_label and participant.name and participant.speaker_label not in mapping:
            mapping[participant.speaker_label] = participant.name
    return mapping


def _format_segment_timestamp(seconds_value: Any) -> str:
    try:
        total_seconds = max(0, int(float(seconds_value or 0)))
    except (TypeError, ValueError):
        total_seconds = 0
    hours, remainder = divmod(total_seconds, 3600)
    minutes, seconds = divmod(remainder, 60)
    if hours:
        return f"{hours:02d}:{minutes:02d}:{seconds:02d}"
    return f"{minutes:02d}:{seconds:02d}"


def _extract_transcript_lines(
    transcript: Optional[models.Transcript],
    meeting: models.Meeting,
    labels: dict[str, str],
    db: Session,
) -> list[str]:
    if not transcript:
        return [labels["no_data"]]
    speaker_lookup = _build_speaker_lookup(meeting)
    segments = (
        db.query(models.TranscriptSegment)
        .filter(models.TranscriptSegment.transcript_id == transcript.id)
        .order_by(models.TranscriptSegment.start_time.asc())
        .all()
    )
    if not segments and transcript.content:
        return [transcript.content.strip()]
    if not segments:
        return [labels["no_data"]]

    lines: list[str] = []
    for segment in segments:
        speaker = speaker_lookup.get(segment.speaker_label, segment.speaker_label or "Speaker")
        timestamp = _format_segment_timestamp(segment.start_time)
        text = (segment.text or "").strip()
        if text:
            lines.append(f"[{timestamp}] {speaker}: {text}")
    return lines or [labels["no_data"]]


def _build_conclusion(
    summary: Optional[models.MeetingSummary],
    decisions: list[str],
    action_items: list[str],
    labels: dict[str, str],
) -> str:
    if summary and summary.meeting_summary and summary.meeting_summary.strip():
        return summary.meeting_summary.strip()
    if decisions:
        return decisions[0]
    if action_items:
        return action_items[0]
    return labels["no_data"]


def _build_minutes_payload(meeting: models.Meeting, request: ExportRequest, db: Session) -> dict[str, Any]:
    language = _normalize_language(request.language)
    labels = _labels(language)
    summary = _pick_summary(meeting.id, language, db) if request.include_summary else None
    transcript = _pick_transcript(meeting.id, language, db)

    summary_text = (
        summary.meeting_summary.strip()
        if summary and summary.meeting_summary and summary.meeting_summary.strip()
        else labels["no_data"]
    )
    key_points = _extract_list(summary.key_points if summary else None) or [labels["no_data"]]
    decisions = _extract_list(summary.decisions if summary else None) or [labels["no_data"]]
    risks = _extract_list(summary.risks if summary else None)
    open_questions = _extract_list(summary.open_questions if summary else None)
    timeline_highlights = _extract_list(summary.timeline_highlights if summary else None)
    speaker_summaries = _extract_list(summary.speaker_summaries if summary else None)

    model_action_items = (
        db.query(models.ActionItem)
        .filter(models.ActionItem.meeting_id == meeting.id)
        .order_by(models.ActionItem.created_at.asc())
        .all()
        if request.include_action_items
        else []
    )
    formatted_action_items = []
    for item in model_action_items:
        deadline = item.due_date.strftime("%d/%m/%Y") if item.due_date else labels["no_data"]
        formatted_action_items.append(
            f"{item.title} | {labels['owner']}: {_format_action_item_assignee_summary(item)}"
            f" | {labels['deadline']}: {deadline}"
            f" | {labels['status']}: {item.status}"
            f" | {labels['priority']}: {item.priority}"
        )

    if not formatted_action_items and summary:
        structured_action_items = _extract_list(summary.action_items)
        formatted_action_items.extend(structured_action_items)

    if not formatted_action_items:
        formatted_action_items = [labels["no_data"]]

    transcript_lines = _extract_transcript_lines(transcript, meeting, labels, db)
    time_parts = _meeting_time_parts(meeting, labels)
    creator = _creator_label(meeting, labels)
    org_name = _clean_value(getattr(meeting.organization, "name", None)) if meeting.organization else ""
    description = _clean_value(meeting.description)
    meeting_type = _meeting_type_label(meeting, labels)

    # Build structured action items for table rendering
    action_items_structured = []
    for idx, item in enumerate(model_action_items, 1):
        action_items_structured.append({
            "no": str(idx),
            "title": item.title or "",
            "owner": _format_action_item_assignee_summary(item),
            "deadline": item.due_date.strftime("%d/%m/%Y") if item.due_date else labels["no_data"],
            "status": item.status or "",
            "priority": item.priority or "",
        })

    return {
        "language": language,
        "labels": labels,
        "organization_name": org_name,
        "meeting_title": _display_meeting_title(meeting, labels),
        "meeting_description": description,
        "meeting_type": meeting_type,
        "meeting_creator": creator,
        "meeting_time": _display_meeting_time(meeting, labels),
        "meeting_time_start": time_parts["start"],
        "meeting_time_end": time_parts["end"],
        "meeting_time_duration": time_parts["duration"],
        "duration": _format_duration(meeting.duration, labels),
        "location_mode": _format_location_or_mode(meeting, labels),
        "participants": _extract_participants(meeting, labels),
        "summary_text": summary_text,
        "key_points": key_points,
        "decisions": decisions,
        "action_items": formatted_action_items,
        "action_items_structured": action_items_structured,
        "risks": risks,
        "open_questions": open_questions,
        "timeline_highlights": timeline_highlights,
        "speaker_summaries": speaker_summaries,
        "conclusion": _build_conclusion(summary, decisions, formatted_action_items, labels),
        "appendix_lines": transcript_lines,
        "include_transcript_appendix": _wants_transcript_appendix(request),
        "generated_at": _format_meeting_datetime(datetime.now(), labels),
        "download_date": _format_meeting_datetime(datetime.now(), labels),
    }


def _resolve_pdf_font_paths(language: str) -> tuple[str, str]:
    candidate_groups = [SERIF_FONT_CANDIDATES, UNICODE_FONT_CANDIDATES] if language in {"vi", "en"} else [UNICODE_FONT_CANDIDATES, SERIF_FONT_CANDIDATES]
    for candidates in candidate_groups:
        for regular_path, bold_path in candidates:
            if Path(regular_path).exists():
                return regular_path, bold_path if Path(bold_path).exists() else regular_path
    raise RuntimeError("No Unicode font available for PDF export")


def _configure_pdf(pdf: FPDF, language: str) -> str:
    regular_path, bold_path = _resolve_pdf_font_paths(language)
    font_family = "MinutesSerif" if language in {"vi", "en"} else "MinutesUnicode"
    pdf.add_font(font_family, "", regular_path)
    pdf.add_font(font_family, "B", bold_path)
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.set_margins(20, 18, 20)
    pdf.add_page()
    return font_family


def _pdf_available_width(pdf: FPDF, inset: float = 0) -> float:
    remaining = pdf.w - pdf.r_margin - pdf.get_x() - inset
    return max(remaining, 30)


def _pdf_safe_multicell(
    pdf: FPDF,
    line_height: float,
    text: str,
    *,
    inset: float = 0,
    align: str = "L",
) -> None:
    width = _pdf_available_width(pdf, inset=inset)
    try:
        pdf.multi_cell(width, line_height, text, align=align)
    except FPDFException:
        # Fall back to per-line rendering so a narrow cursor state does not kill the export.
        for raw_line in (text or "").splitlines() or [""]:
            pdf.multi_cell(width, line_height, raw_line or " ", align=align)


def _pdf_section_title(pdf: FPDF, font_family: str, text: str, index: int) -> None:
    pdf.set_font(font_family, "B", 13)
    pdf.set_text_color(0, 0, 0)
    pdf.ln(3)
    pdf.cell(0, 8, f"{index}. {text}", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(1)


def _pdf_key_value(pdf: FPDF, font_family: str, label: str, value: str) -> None:
    pdf.set_text_color(0, 0, 0)
    pdf.set_font(font_family, "B", 11)
    label_w = pdf.get_string_width(f"{label}: ")
    value_w = pdf.get_string_width(value)
    total_w = label_w + value_w
    avail = pdf.w - pdf.l_margin - pdf.r_margin

    if total_w <= avail:
        # Fits on one line: bold label + normal value
        pdf.cell(label_w + 2, 7, f"{label}: ", new_x="RIGHT", new_y="TOP")
        pdf.set_font(font_family, "", 11)
        pdf.cell(0, 7, value, new_x="LMARGIN", new_y="NEXT")
    else:
        # Label on its own line, value indented below
        pdf.set_font(font_family, "B", 11)
        _pdf_safe_multicell(pdf, 7, f"{label}:")
        pdf.set_font(font_family, "", 11)
        pdf.set_x(pdf.l_margin + 6)
        _pdf_safe_multicell(pdf, 7, value)


def _pdf_bullet_list(pdf: FPDF, font_family: str, items: list[str]) -> None:
    for item in items:
        pdf.set_font(font_family, "", 11)
        pdf.set_x(pdf.l_margin + 2)
        _pdf_safe_multicell(pdf, 7, f"- {item}")
        pdf.ln(0.5)


def _pdf_action_items_table(pdf: FPDF, font_family: str, structured_items: list[dict], labels: dict[str, str]) -> None:
    if not structured_items:
        pdf.set_font(font_family, "", 11)
        _pdf_safe_multicell(pdf, 7, labels["no_data"])
        return

    col_widths = [12, 55, 35, 25, 25, 20]  # STT, Nội dung, Phụ trách, Hạn, Trạng thái, Ưu tiên
    headers = [
        labels["action_table_no"],
        labels["action_table_title"],
        labels["action_table_owner"],
        labels["action_table_deadline"],
        labels["action_table_status"],
        labels["action_table_priority"],
    ]
    row_height = 7

    # Header row
    pdf.set_font(font_family, "B", 9)
    pdf.set_fill_color(240, 240, 240)
    for i, header in enumerate(headers):
        pdf.cell(col_widths[i], row_height, header, border=1, fill=True, align="C")
    pdf.ln(row_height)

    # Data rows
    pdf.set_font(font_family, "", 9)
    for item in structured_items:
        values = [
            item["no"],
            item["title"],
            item["owner"],
            item["deadline"],
            item["status"],
            item["priority"],
        ]
        # Calculate max height needed for this row
        max_lines = 1
        for i, val in enumerate(values):
            text_width = pdf.get_string_width(val)
            lines = max(1, int(text_width / (col_widths[i] - 2)) + 1)
            max_lines = max(max_lines, lines)
        cell_height = row_height * max_lines

        for i, val in enumerate(values):
            align = "C" if i in (0, 3, 4, 5) else "L"
            pdf.cell(col_widths[i], cell_height, val[:40] if i == 2 else val[:60], border=1, align=align)
        pdf.ln(cell_height)


def _pdf_signature_block(pdf: FPDF, font_family: str, labels: dict[str, str], download_date: str = "") -> None:
    pdf.ln(6)
    # Date line above signatures (Vietnamese formal format)
    if download_date:
        pdf.set_font(font_family, "", 11)
        pdf.set_text_color(0, 0, 0)
        pdf.cell(0, 7, download_date, align="R", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4)

    column_width = (pdf.w - pdf.l_margin - pdf.r_margin) / 3
    start_y = pdf.get_y()
    titles = [labels["chairperson"], labels["secretary"], labels["confirmation"]]

    for index, title in enumerate(titles):
        x = pdf.l_margin + index * column_width
        pdf.set_xy(x, start_y)
        pdf.set_font(font_family, "B", 11)
        pdf.cell(column_width, 7, title, align="C")
        pdf.line(x + 12, start_y + 38, x + column_width - 12, start_y + 38)
        pdf.set_xy(x, start_y + 40)
        pdf.set_font(font_family, "", 9)
        pdf.cell(column_width, 6, labels["signature_placeholder"], align="C")

    pdf.set_y(start_y + 52)


def create_pdf_export(meeting: models.Meeting, request: ExportRequest, db: Session) -> bytes:
    payload = _build_minutes_payload(meeting, request, db)
    labels = payload["labels"]

    pdf = FPDF()
    font_family = _configure_pdf(pdf, payload["language"])

    # Header: Organization name (right-aligned)
    if payload["organization_name"]:
        pdf.set_font(font_family, "", 11)
        pdf.set_text_color(80, 80, 80)
        pdf.cell(0, 7, f"{labels['organization_name']}: {payload['organization_name']}", align="R", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

    # Document title
    pdf.set_font(font_family, "B", 18)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(0, 10, labels["document_title"], align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font(font_family, "", 11)
    pdf.cell(0, 6, labels["document_subtitle"], align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(6)

    # Section 1: Meeting info
    section_index = 1
    _pdf_section_title(pdf, font_family, labels["main_content"], section_index)
    _pdf_key_value(pdf, font_family, labels["meeting_name"], payload["meeting_title"])
    if payload["meeting_description"]:
        _pdf_key_value(pdf, font_family, labels["meeting_description"], payload["meeting_description"])
    _pdf_key_value(pdf, font_family, labels["time_start"], payload["meeting_time_start"])
    _pdf_key_value(pdf, font_family, labels["time_end"], payload["meeting_time_end"])
    _pdf_key_value(pdf, font_family, labels["duration"], payload["meeting_time_duration"])
    _pdf_key_value(pdf, font_family, labels["location_mode"], payload["location_mode"])
    if payload["meeting_creator"]:
        _pdf_key_value(pdf, font_family, labels["meeting_creator"], payload["meeting_creator"])

    # Section 2: Participants
    section_index += 1
    _pdf_section_title(pdf, font_family, labels["participants"], section_index)
    _pdf_bullet_list(pdf, font_family, payload["participants"])

    # Section 3: Summary
    section_index += 1
    _pdf_section_title(pdf, font_family, labels["summary"], section_index)
    pdf.set_font(font_family, "", 11)
    _pdf_safe_multicell(pdf, 7, payload["summary_text"])

    # Section 4: Key points
    section_index += 1
    _pdf_section_title(pdf, font_family, labels["key_points"], section_index)
    _pdf_bullet_list(pdf, font_family, payload["key_points"])

    # Section 5: Decisions
    section_index += 1
    _pdf_section_title(pdf, font_family, labels["decisions"], section_index)
    _pdf_bullet_list(pdf, font_family, payload["decisions"])

    # Section 6: Risks (if data exists)
    if payload["risks"]:
        section_index += 1
        _pdf_section_title(pdf, font_family, labels["risks"], section_index)
        _pdf_bullet_list(pdf, font_family, payload["risks"])

    # Section 7: Open questions (if data exists)
    if payload["open_questions"]:
        section_index += 1
        _pdf_section_title(pdf, font_family, labels["open_questions"], section_index)
        _pdf_bullet_list(pdf, font_family, payload["open_questions"])

    # Section 8: Timeline highlights (if data exists)
    if payload["timeline_highlights"]:
        section_index += 1
        _pdf_section_title(pdf, font_family, labels["timeline_highlights"], section_index)
        _pdf_bullet_list(pdf, font_family, payload["timeline_highlights"])

    # Section 9: Speaker summaries (if data exists)
    if payload["speaker_summaries"]:
        section_index += 1
        _pdf_section_title(pdf, font_family, labels["speaker_summaries"], section_index)
        _pdf_bullet_list(pdf, font_family, payload["speaker_summaries"])

    # Section 10: Action items as table
    section_index += 1
    _pdf_section_title(pdf, font_family, labels["action_items"], section_index)
    if payload["action_items_structured"]:
        _pdf_action_items_table(pdf, font_family, payload["action_items_structured"], labels)
    else:
        _pdf_bullet_list(pdf, font_family, payload["action_items"])

    # Section 11: Conclusion
    section_index += 1
    _pdf_section_title(pdf, font_family, labels["conclusion"], section_index)
    pdf.set_font(font_family, "", 11)
    _pdf_safe_multicell(pdf, 7, payload["conclusion"])

    # Section 12: Signature
    section_index += 1
    _pdf_section_title(pdf, font_family, labels["signature_section"], section_index)
    _pdf_signature_block(pdf, font_family, labels, payload["download_date"])

    # Appendix
    if payload["include_transcript_appendix"]:
        pdf.add_page()
        pdf.set_font(font_family, "B", 14)
        pdf.cell(0, 8, labels["appendix_title"], new_x="LMARGIN", new_y="NEXT")
        pdf.set_font(font_family, "", 10)
        pdf.set_text_color(107, 114, 128)
        _pdf_safe_multicell(pdf, 6, labels["appendix_note"])
        pdf.ln(2)
        pdf.set_text_color(17, 24, 39)
        for line in payload["appendix_lines"]:
            _pdf_safe_multicell(pdf, 6, line)
            pdf.ln(1)

    output = pdf.output(dest="S")
    if isinstance(output, str):
        return output.encode("latin-1", "replace")
    return bytes(output)


def _apply_docx_base_style(document: Document) -> None:
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)
    document.sections[0].top_margin = Inches(0.85)
    document.sections[0].bottom_margin = Inches(0.85)
    document.sections[0].left_margin = Inches(1.0)
    document.sections[0].right_margin = Inches(0.85)


def _docx_add_key_value(document: Document, label: str, value: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.2
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    paragraph.add_run(value)


def _docx_add_bullet_list(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def _docx_action_items_table(document: Document, structured_items: list[dict], labels: dict[str, str]) -> None:
    if not structured_items:
        document.add_paragraph(labels["no_data"])
        return

    headers = [
        labels["action_table_no"],
        labels["action_table_title"],
        labels["action_table_owner"],
        labels["action_table_deadline"],
        labels["action_table_status"],
        labels["action_table_priority"],
    ]
    table = document.add_table(rows=1 + len(structured_items), cols=6)
    table.style = "Table Grid"

    # Header row
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Data rows
    for row_idx, item in enumerate(structured_items, 1):
        values = [item["no"], item["title"], item["owner"], item["deadline"], item["status"], item["priority"]]
        for col_idx, val in enumerate(values):
            cell = table.cell(row_idx, col_idx)
            cell.text = val
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx in (0, 3, 4, 5) else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(10)


def _docx_add_signature_table(document: Document, labels: dict[str, str], download_date: str = "") -> None:
    # Date line above signatures
    if download_date:
        date_para = document.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_para.paragraph_format.space_after = Pt(12)
        run = date_para.add_run(download_date)
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"

    table = document.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    titles = [labels["chairperson"], labels["secretary"], labels["confirmation"]]
    for index, title in enumerate(titles):
        header = table.cell(0, index)
        header.text = title
        header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(1, index)
        cell.text = f"\n\n\n{labels['signature_placeholder']}"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def create_docx_export(meeting: models.Meeting, request: ExportRequest, db: Session) -> bytes:
    payload = _build_minutes_payload(meeting, request, db)
    labels = payload["labels"]

    document = Document()
    _apply_docx_base_style(document)

    # Header: Organization name (right-aligned)
    if payload["organization_name"]:
        org_para = document.add_paragraph()
        org_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        org_para.paragraph_format.space_after = Pt(2)
        run = org_para.add_run(f"{labels['organization_name']}: {payload['organization_name']}")
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"

    # Document title
    title = document.add_heading(labels["document_title"], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.name = "Times New Roman"
    subtitle = document.add_paragraph(labels["document_subtitle"])
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True
    subtitle.runs[0].font.name = "Times New Roman"

    # Section 1: Meeting info
    section_index = 1
    document.add_heading(f"{section_index}. {labels['main_content']}", level=1)
    _docx_add_key_value(document, labels["meeting_name"], payload["meeting_title"])
    if payload["meeting_description"]:
        _docx_add_key_value(document, labels["meeting_description"], payload["meeting_description"])
    _docx_add_key_value(document, labels["time_start"], payload["meeting_time_start"])
    _docx_add_key_value(document, labels["time_end"], payload["meeting_time_end"])
    _docx_add_key_value(document, labels["duration"], payload["meeting_time_duration"])
    _docx_add_key_value(document, labels["location_mode"], payload["location_mode"])
    if payload["meeting_creator"]:
        _docx_add_key_value(document, labels["meeting_creator"], payload["meeting_creator"])

    # Section 2: Participants
    section_index += 1
    document.add_heading(f"{section_index}. {labels['participants']}", level=1)
    _docx_add_bullet_list(document, payload["participants"])

    # Section 3: Summary
    section_index += 1
    document.add_heading(f"{section_index}. {labels['summary']}", level=1)
    document.add_paragraph(payload["summary_text"])

    # Section 4: Key points
    section_index += 1
    document.add_heading(f"{section_index}. {labels['key_points']}", level=1)
    _docx_add_bullet_list(document, payload["key_points"])

    # Section 5: Decisions
    section_index += 1
    document.add_heading(f"{section_index}. {labels['decisions']}", level=1)
    _docx_add_bullet_list(document, payload["decisions"])

    # Section 6: Risks (if data exists)
    if payload["risks"]:
        section_index += 1
        document.add_heading(f"{section_index}. {labels['risks']}", level=1)
        _docx_add_bullet_list(document, payload["risks"])

    # Section 7: Open questions (if data exists)
    if payload["open_questions"]:
        section_index += 1
        document.add_heading(f"{section_index}. {labels['open_questions']}", level=1)
        _docx_add_bullet_list(document, payload["open_questions"])

    # Section 8: Timeline highlights (if data exists)
    if payload["timeline_highlights"]:
        section_index += 1
        document.add_heading(f"{section_index}. {labels['timeline_highlights']}", level=1)
        _docx_add_bullet_list(document, payload["timeline_highlights"])

    # Section 9: Speaker summaries (if data exists)
    if payload["speaker_summaries"]:
        section_index += 1
        document.add_heading(f"{section_index}. {labels['speaker_summaries']}", level=1)
        _docx_add_bullet_list(document, payload["speaker_summaries"])

    # Section 10: Action items as table
    section_index += 1
    document.add_heading(f"{section_index}. {labels['action_items']}", level=1)
    if payload["action_items_structured"]:
        _docx_action_items_table(document, payload["action_items_structured"], labels)
    else:
        _docx_add_bullet_list(document, payload["action_items"])

    # Section 11: Conclusion
    section_index += 1
    document.add_heading(f"{section_index}. {labels['conclusion']}", level=1)
    document.add_paragraph(payload["conclusion"])

    # Section 12: Signature with download date
    section_index += 1
    document.add_paragraph()
    document.add_heading(f"{section_index}. {labels['signature_section']}", level=1)
    _docx_add_signature_table(document, labels, payload["download_date"])

    # Appendix
    if payload["include_transcript_appendix"]:
        document.add_page_break()
        document.add_heading(labels["appendix_title"], level=1)
        note = document.add_paragraph(labels["appendix_note"])
        note.runs[0].italic = True
        for line in payload["appendix_lines"]:
            document.add_paragraph(line)

    doc_stream = io.BytesIO()
    document.save(doc_stream)
    return doc_stream.getvalue()


@router.post("/generate", response_model=ExportResponse)
async def generate_export(
    request: ExportRequest,
    db: Session = Depends(get_db),
    current_user=Depends(auth.get_current_user),
):
    meeting = db.query(models.Meeting).filter(models.Meeting.id == request.meeting_id).first()
    if not meeting:
        raise HTTPException(status_code=404, detail="Meeting not found")
    auth.require_org_member(db, current_user, meeting.organization_id)

    try:
        export_format = request.format.lower()
        if export_format == "pdf":
            file_content = create_pdf_export(meeting, request, db)
            extension = "pdf"
        elif export_format == "docx":
            file_content = create_docx_export(meeting, request, db)
            extension = "docx"
        else:
            raise HTTPException(status_code=400, detail="Unsupported format")

        title_slug = _filename_title_slug(meeting)
        group_name = _clean_value(getattr(meeting.group, "name", None)) if meeting.group else ""
        org_name = _clean_value(getattr(meeting.organization, "name", None)) if meeting.organization else ""
        location_slug = _slugify_filename_part(group_name or org_name or "meeting")
        start_time = meeting.actual_start or meeting.scheduled_start or meeting.created_at
        time_slug = start_time.strftime("%d-%m-%Y_%H-%M") if start_time else datetime.now().strftime("%d-%m-%Y_%H-%M")
        filename = f"{title_slug}_{location_slug}_{time_slug}.{extension}"
        file_path = os.path.join(tempfile.gettempdir(), filename)
        with open(file_path, "wb") as handle:
            handle.write(file_content)

        return ExportResponse(
            download_url=f"/api/export/download/{filename}?meeting_id={meeting.id}",
            filename=filename,
            size_bytes=len(file_content),
            created_at=datetime.now(),
        )
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Error generating export: {exc}") from exc


@router.get("/download/{filename}")
async def download_export(
    filename: str,
    meeting_id: Optional[str] = None,
    db: Session = Depends(get_db),
    current_user=Depends(auth.get_current_user),
):
    if not re.match(r"^[a-zA-Z0-9_\-\.]+$", filename) or ".." in filename or "/" in filename:
        raise HTTPException(status_code=400, detail="Invalid filename")

    filename_match = EXPORT_FILENAME_PATTERN.fullmatch(filename)
    if not filename_match:
        raise HTTPException(status_code=400, detail="Invalid filename")

    file_path = os.path.join(tempfile.gettempdir(), filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")

    # Try meeting_id from query param first, then from filename
    resolved_meeting_id = meeting_id
    if not resolved_meeting_id:
        raise HTTPException(status_code=400, detail="meeting_id is required")

    meeting = db.query(models.Meeting).filter(models.Meeting.id == resolved_meeting_id).first()
    if not meeting:
        raise HTTPException(status_code=404, detail="Meeting not found")
    auth.require_org_member(db, current_user, meeting.organization_id)

    if filename.endswith(".pdf"):
        content_type = "application/pdf"
    elif filename.endswith(".docx"):
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        content_type = "application/octet-stream"

    with open(file_path, "rb") as handle:
        file_content = handle.read()

    return Response(
        content=file_content,
        media_type=content_type,
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )
