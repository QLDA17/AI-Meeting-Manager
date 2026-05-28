import pytest
import json
from unittest.mock import patch
from io import BytesIO
from pathlib import Path
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from sqlalchemy.pool import StaticPool
from docx import Document
from src.api.main import app
from src.api.database import get_db, Base
from src.api import models
from src.api.crud import add_user_to_organization, create_organization, create_user

TEST_DATABASE_URL = "sqlite:///:memory:"
test_engine = create_engine(
    TEST_DATABASE_URL,
    connect_args={"check_same_thread": False},
    poolclass=StaticPool,
)
TestingSessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=test_engine)


def override_get_db():
    db = TestingSessionLocal()
    canonical_summary_id = None
    try:
        yield db
    finally:
        db.close()


app.dependency_overrides[get_db] = override_get_db


def activate_org(org_id: str):
    db = TestingSessionLocal()
    try:
        org = db.query(models.Organization).filter(models.Organization.id == org_id).first()
        org.settings = {"approval_status": "active"}
        db.commit()
    finally:
        db.close()


@pytest.fixture(scope="function")
def client():
    Base.metadata.create_all(bind=test_engine)
    app.dependency_overrides[get_db] = override_get_db
    with TestClient(app) as test_client:
        yield test_client
    app.dependency_overrides.pop(get_db, None)
    Base.metadata.drop_all(bind=test_engine)


@pytest.fixture
def auth_context(client):
    db = TestingSessionLocal()
    try:
        user = create_user(
            db,
            {
                "username": "testuser",
                "email": "test@example.com",
                "password": "securepassword123",
                "role": "member",
            },
        )
        org = create_organization(db, {"name": "Test Organization", "settings": {"approval_status": "active"}})
        add_user_to_organization(db, user.id, org.id, "org-admin")
    finally:
        db.close()

    response = client.post(
        "/api/auth/login",
        json={
            "username": "testuser",
            "password": "securepassword123",
        },
    )
    data = response.json()
    return {
        "headers": {"Authorization": f"Bearer {data['access_token']}"},
        "org_id": data["user"]["orgMemberships"][0]["orgId"],
    }


def meeting_payload(auth_context, **overrides):
    data = {
        "organization_id": auth_context["org_id"],
        "title": "Test Meeting",
        "description": "A test meeting",
        "meeting_type": "MEETING",
        "status": "upcoming",
        "scheduled_start": "2099-05-27T09:00:00Z",
        "scheduled_end": "2099-05-27T10:00:00Z",
    }
    data.update(overrides)
    return data


def test_create_meeting(client, auth_context):
    response = client.post(
        "/api/meetings",
        headers=auth_context["headers"],
        json=meeting_payload(auth_context),
    )

    assert response.status_code == 200
    data = response.json()
    assert data["title"] == "Test Meeting"
    assert data["status"] == "upcoming"
    assert data["meeting_type"] == "MEETING"
    assert data["organization_id"] == auth_context["org_id"]
    assert "id" in data


def test_create_upcoming_meeting_requires_scheduled_start(client, auth_context):
    response = client.post(
        "/api/meetings",
        headers=auth_context["headers"],
        json=meeting_payload(auth_context, status="upcoming", scheduled_start=None, scheduled_end=None),
    )

    assert response.status_code == 400
    assert "thời gian bắt đầu" in response.json()["detail"].lower()


def test_create_meeting_unauthorized(client):
    response = client.post(
        "/api/meetings",
        json={"title": "Test Meeting"},
    )

    assert response.status_code == 401


def test_get_meetings(client, auth_context):
    first_response = client.post(
        "/api/meetings",
        headers=auth_context["headers"],
        json=meeting_payload(
            auth_context,
            title="Test Meeting 1",
            scheduled_start="2099-05-27T09:00:00Z",
            scheduled_end="2099-05-27T10:00:00Z",
        ),
    )
    second_response = client.post(
        "/api/meetings",
        headers=auth_context["headers"],
        json=meeting_payload(
            auth_context,
            title="Test Meeting 2",
            scheduled_start="2099-05-27T11:00:00Z",
            scheduled_end="2099-05-27T12:00:00Z",
        ),
    )

    assert first_response.status_code == 200
    assert second_response.status_code == 200

    response = client.get("/api/meetings", headers=auth_context["headers"])

    assert response.status_code == 200
    data = response.json()
    assert isinstance(data, list)
    assert len(data) >= 2


def test_list_meetings_exposes_planned_duration_without_overwriting_actual_duration(client, auth_context):
    create_response = client.post(
        "/api/meetings",
        headers=auth_context["headers"],
        json=meeting_payload(
            auth_context,
            title="Scheduled Duration Split",
            scheduled_start="2099-05-27T09:00:00Z",
            scheduled_end="2099-05-27T10:00:00Z",
            duration=0,
        ),
    )
    assert create_response.status_code == 200
    meeting_id = create_response.json()["id"]

    response = client.get("/api/meetings", headers=auth_context["headers"])

    assert response.status_code == 200
    row = next(item for item in response.json() if item["id"] == meeting_id)
    assert row["duration"] == 0
    assert row["planned_duration_minutes"] == 60
    assert row["actual_duration_minutes"] is None
    assert row["live_duration_minutes"] is None
    assert row["is_overrun"] is False


def test_get_meeting_by_id(client, auth_context):
    create_response = client.post(
        "/api/meetings",
        headers=auth_context["headers"],
        json=meeting_payload(auth_context),
    )
    meeting_id = create_response.json()["id"]

    response = client.get(f"/api/meetings/{meeting_id}", headers=auth_context["headers"])

    assert response.status_code == 200
    data = response.json()
    assert data["id"] == meeting_id
    assert data["title"] == "Test Meeting"


def test_get_meeting_detail_exposes_duration_modes(client, auth_context):
    create_response = client.post(
        "/api/meetings",
        headers=auth_context["headers"],
        json=meeting_payload(
            auth_context,
            title="Detail Duration Split",
            status="live",
            scheduled_start="2099-05-27T09:00:00Z",
            scheduled_end="2099-05-27T10:00:00Z",
            actual_start="2099-05-27T09:05:00Z",
        ),
    )
    meeting_id = create_response.json()["id"]

    response = client.get(f"/api/meetings/{meeting_id}", headers=auth_context["headers"])

    assert response.status_code == 200
    data = response.json()
    assert data["duration"] == 0
    assert data["planned_duration_minutes"] == 60
    assert data["actual_duration_minutes"] is None
    assert data["live_duration_minutes"] is not None
    assert data["is_overrun"] is False


def test_get_meeting_not_found(client, auth_context):
    response = client.get("/api/meetings/nonexistent-id", headers=auth_context["headers"])

    assert response.status_code == 404
    assert "not found" in response.json()["detail"].lower()


def test_update_meeting(client, auth_context):
    create_response = client.post(
        "/api/meetings",
        headers=auth_context["headers"],
        json=meeting_payload(auth_context, title="Original Title"),
    )
    meeting_id = create_response.json()["id"]

    response = client.put(
        f"/api/meetings/{meeting_id}",
        headers=auth_context["headers"],
        json={
            "title": "Updated Title",
            "status": "live",
            "description": "Updated description",
        },
    )

    assert response.status_code == 200
    data = response.json()
    assert data["title"] == "Updated Title"
    assert data["status"] == "live"
    assert data["description"] == "Updated description"


def test_update_meeting_rejects_past_scheduled_start(client, auth_context):
    create_response = client.post(
        "/api/meetings",
        headers=auth_context["headers"],
        json=meeting_payload(
            auth_context,
            title="Future Meeting",
            scheduled_start="2026-05-27T09:00:00Z",
            scheduled_end="2026-05-27T10:00:00Z",
        ),
    )
    meeting_id = create_response.json()["id"]

    response = client.put(
        f"/api/meetings/{meeting_id}",
        headers=auth_context["headers"],
        json={"scheduled_start": "2026-05-20T09:00:00Z"},
    )

    assert response.status_code == 400
    assert "quá khứ" in response.json()["detail"].lower()


def test_update_meeting_rejects_completed_meeting(client, auth_context):
    create_response = client.post(
        "/api/meetings",
        headers=auth_context["headers"],
        json=meeting_payload(
            auth_context,
            title="Completed Meeting",
            scheduled_start="2026-05-27T09:00:00Z",
            scheduled_end="2026-05-27T10:00:00Z",
        ),
    )
    meeting_id = create_response.json()["id"]

    end_response = client.post(
        f"/api/meetings/{meeting_id}/end",
        headers=auth_context["headers"],
        json={"status": "completed"},
    )
    assert end_response.status_code == 200

    response = client.put(
        f"/api/meetings/{meeting_id}",
        headers=auth_context["headers"],
        json={"title": "Should be blocked"},
    )

    assert response.status_code == 400
    assert "đã kết thúc" in response.json()["detail"].lower()


def test_delete_meeting(client, auth_context):
    create_response = client.post(
        "/api/meetings",
        headers=auth_context["headers"],
        json=meeting_payload(auth_context),
    )
    meeting_id = create_response.json()["id"]

    response = client.delete(f"/api/meetings/{meeting_id}", headers=auth_context["headers"])

    assert response.status_code == 200
    assert "deleted successfully" in response.json()["message"].lower()

    get_response = client.get(f"/api/meetings/{meeting_id}", headers=auth_context["headers"])
    assert get_response.status_code == 404


def test_get_meetings_with_filters(client, auth_context):
    client.post(
        "/api/meetings",
        headers=auth_context["headers"],
        json=meeting_payload(auth_context, title="Upcoming Meeting", status="upcoming"),
    )
    client.post(
        "/api/meetings",
        headers=auth_context["headers"],
        json=meeting_payload(auth_context, title="Completed Meeting", status="completed"),
    )

    response = client.get("/api/meetings?status=upcoming", headers=auth_context["headers"])

    assert response.status_code == 200
    data = response.json()
    assert all(m["status"] == "upcoming" for m in data)


def test_get_meetings_sorted_by_meeting_time_desc(client, auth_context):
    older_response = client.post(
        "/api/meetings",
        headers=auth_context["headers"],
        json=meeting_payload(
            auth_context,
            title="Older Start",
            scheduled_start="2026-05-27T09:00:00Z",
            scheduled_end="2026-05-27T10:00:00Z",
        ),
    )
    newer_response = client.post(
        "/api/meetings",
        headers=auth_context["headers"],
        json=meeting_payload(
            auth_context,
            title="Newer Start",
            scheduled_start="2026-05-28T09:00:00Z",
            scheduled_end="2026-05-28T10:00:00Z",
        ),
    )

    assert older_response.status_code == 200
    assert newer_response.status_code == 200

    response = client.get("/api/meetings", headers=auth_context["headers"])

    assert response.status_code == 200
    data = response.json()
    titles = [meeting["title"] for meeting in data[:2]]
    assert titles == ["Newer Start", "Older Start"]


def test_upload_audio_creates_meeting_audio_and_job(client, auth_context, monkeypatch, tmp_path):
    from src.api.routes import stt as stt_routes

    monkeypatch.setattr(stt_routes, "ensure_audio_upload_dir", lambda: str(tmp_path))
    monkeypatch.setattr(stt_routes, "start_upload_job", lambda job: None)

    response = client.post(
        "/api/upload",
        headers=auth_context["headers"],
        files={"file": ("backup.wav", b"fake-wave-data", "audio/wav")},
        data={
            "organization_id": auth_context["org_id"],
            "title": "Backup Upload",
            "language": "auto",
            "stt_provider": "deepgram",
            "enable_diarization": "true",
            "enable_summary": "true",
            "enable_noise_cleanup": "true",
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["status"] == "queued"
    assert payload["current_stage"] == "queued"
    assert payload["progress_percent"] == 0

    status_response = client.get(f"/api/upload/jobs/{payload['job_id']}", headers=auth_context["headers"])
    assert status_response.status_code == 200
    assert status_response.json()["meeting_id"] == payload["meeting_id"]

    db = TestingSessionLocal()
    try:
        meeting = db.query(models.Meeting).filter(models.Meeting.id == payload["meeting_id"]).first()
        assert meeting is not None
        assert meeting.status == "queued"
        assert meeting.audio_url
        assert meeting.recording_url

        audio_file = db.query(models.AudioFile).filter(models.AudioFile.meeting_id == meeting.id).first()
        assert audio_file is not None
        assert audio_file.upload_status == "UPLOADED"
        assert Path(audio_file.file_path).exists()
        from src.api.core.upload_jobs import get_upload_job
        job = get_upload_job(payload["job_id"])
        assert job is not None
        assert job.enable_action_items is False
    finally:
        db.close()


def test_upload_audio_batch_creates_one_meeting_per_file(client, auth_context, monkeypatch, tmp_path):
    from src.api.routes import stt as stt_routes

    monkeypatch.setattr(stt_routes, "ensure_audio_upload_dir", lambda: str(tmp_path))
    monkeypatch.setattr(stt_routes, "start_upload_job", lambda job: None)

    response = client.post(
        "/api/uploads/batch",
        headers=auth_context["headers"],
        files=[
            ("files", ("first.wav", b"fake-wave-data-1", "audio/wav")),
            ("files", ("second.wav", b"fake-wave-data-2", "audio/wav")),
        ],
        data={
            "organization_id": auth_context["org_id"],
            "items": json.dumps([
                {
                    "client_id": "client-1",
                    "title": "First Upload",
                    "language": "vi",
                    "stt_provider": "deepgram",
                    "group_id": "",
                },
                {
                    "client_id": "client-2",
                    "title": "Second Upload",
                    "language": "auto",
                    "stt_provider": "viwhisper",
                    "group_id": "",
                },
            ]),
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["total_items"] == 2
    assert payload["status"] == "queued"
    assert len(payload["items"]) == 2
    assert payload["items"][0]["client_id"] == "client-1"
    assert payload["items"][1]["client_id"] == "client-2"

    batch_status = client.get(f"/api/uploads/batch/{payload['batch_id']}", headers=auth_context["headers"])
    assert batch_status.status_code == 200
    assert batch_status.json()["total_items"] == 2

    db = TestingSessionLocal()
    try:
        meetings = db.query(models.Meeting).filter(models.Meeting.title.in_(["First Upload", "Second Upload"])).all()
        assert len(meetings) == 2
        audio_files = db.query(models.AudioFile).filter(models.AudioFile.meeting_id.in_([meeting.id for meeting in meetings])).all()
        assert len(audio_files) == 2
        assert all(Path(audio_file.file_path).exists() for audio_file in audio_files)
    finally:
        db.close()


def test_get_meetings_unauthorized(client):
    response = client.get("/api/meetings")

    assert response.status_code == 401


def test_finalize_meeting_single_router_call_success(client, auth_context):
    create_response = client.post(
        "/api/meetings",
        headers=auth_context["headers"],
        json=meeting_payload(auth_context, title="Finalize Success", status="processing"),
    )
    meeting_id = create_response.json()["id"]

    router_response = """
    {
      "meeting_summary": "Cuoc hop thong nhat chot scope STT va summary.",
      "key_points": ["Chot luong live chunk", "Dung Router cho summary"],
      "decisions": ["Bo 4 call rieng trong finalize"],
      "action_items": [{"task": "Cap nhat finalize", "owner": "Backend", "deadline": "2026-05-20"}]
    }
    """

    with patch.dict(
        "os.environ",
        {"ROUTER_API_URL": "http://router.test", "ROUTER_API_KEY": "secret", "ROUTER_MODEL": "router-model"},
        clear=False,
    ), patch(
        "src.providers.router_llm.RouterLLMAdapter.structured_completion",
        return_value=router_response,
    ) as structured_completion:
        finalize_response = client.post(
            f"/api/meetings/{meeting_id}/finalize",
            headers=auth_context["headers"],
            json={
                "transcript": "Hom nay chung ta chot luong live chunk va dung Router de tom tat.",
                "segments": [{"speaker": "Speaker_01", "start": 0, "end": 5, "text": "Noi dung hop"}],
                "full_regenerate": False,
            },
        )

    assert finalize_response.status_code == 200
    data = finalize_response.json()
    assert data["meeting_id"] == meeting_id
    assert data["transcript_status"] == "COMPLETED"
    assert data["summary_status"] == "COMPLETED"
    assert data["summary"]["meeting_summary"] == "Cuoc hop thong nhat chot scope STT va summary."
    assert data["summary"]["key_points"] == ["Chot luong live chunk", "Dung Router cho summary"]
    assert len(data["summary"]["action_items"]) == 1
    assert data["errors"] == []
    assert structured_completion.call_count >= 1

    detail_response = client.get(f"/api/meetings/{meeting_id}", headers=auth_context["headers"])
    assert detail_response.status_code == 200
    detail = detail_response.json()
    assert detail["transcript_content"] == "Hom nay chung ta chot luong live chunk va dung Router de tom tat."
    assert detail["cleaned_transcript_content"] == "Hom nay chung ta chot luong live chunk va dung Router de tom tat."
    assert detail["raw_transcript_content"] == "Hom nay chung ta chot luong live chunk va dung Router de tom tat."
    assert detail["cleaned_transcript_segments"][0]["text"] == "Noi dung hop"
    assert detail["raw_transcript_segments"][0]["text"] == "Noi dung hop"
    assert detail["transcript_quality_metadata"]["provider"] == "deepgram"
    assert detail["meeting_summary_text"] == "Cuoc hop thong nhat chot scope STT va summary."
    assert detail["key_points_text"] == ["Chot luong live chunk", "Dung Router cho summary"]
    assert detail["decisions_text"] == ["Bo 4 call rieng trong finalize"]
    assert detail["action_items"] == []


def test_finalize_meeting_router_failure_marks_summary_failed(client, auth_context):
    create_response = client.post(
        "/api/meetings",
        headers=auth_context["headers"],
        json=meeting_payload(auth_context, title="Finalize Failure", status="processing"),
    )
    meeting_id = create_response.json()["id"]

    with patch.dict(
        "os.environ",
        {"ROUTER_API_URL": "http://router.test", "ROUTER_API_KEY": "secret", "ROUTER_MODEL": "router-model"},
        clear=False,
    ), patch(
        "src.providers.router_llm.RouterLLMAdapter.structured_completion",
        return_value="not valid json",
    ):
        finalize_response = client.post(
            f"/api/meetings/{meeting_id}/finalize",
            headers=auth_context["headers"],
            json={
                "transcript": "Transcript van duoc luu du Router tra ve sai dinh dang.",
                "segments": [],
                "full_regenerate": False,
            },
        )

    assert finalize_response.status_code == 200
    data = finalize_response.json()
    assert data["transcript_status"] == "COMPLETED"
    assert data["summary_status"] == "FAILED"
    assert data["summary"]["meeting_summary"] == ""
    assert data["summary"]["key_points"] == []
    assert data["summary"]["decisions"] == []
    assert data["errors"]

    detail_response = client.get(f"/api/meetings/{meeting_id}", headers=auth_context["headers"])
    assert detail_response.status_code == 200
    detail = detail_response.json()
    assert detail["transcript_content"] == "Transcript van duoc luu du Router tra ve sai dinh dang."
    assert detail["meeting_summary_text"] is None
    assert detail["key_points_text"] == []
    assert detail["decisions_text"] == []
    assert detail["action_items"] == []


def test_meeting_detail_returns_raw_and_cleaned_transcript_variants(client, auth_context):
    db = TestingSessionLocal()
    try:
        user = db.query(models.User).filter(models.User.email == "test@example.com").first()
        meeting = models.Meeting(
            organization_id=auth_context["org_id"],
            title="Transcript Variants",
            description="Raw and cleaned payload",
            meeting_type="MEETING",
            status="completed",
            created_by=user.id,
        )
        db.add(meeting)
        db.commit()
        db.refresh(meeting)

        transcript = models.Transcript(
            meeting_id=meeting.id,
            content="KPI da duoc chot.",
            raw_content="ka pi ai da duoc chot",
            language="vi",
            word_count=5,
            processing_status="COMPLETED",
            stt_provider="deepgram",
            post_processed=True,
            nlp_metadata={
                "quality_metadata": {
                    "provider": "deepgram",
                    "provider_model": "nova-3",
                    "detected_language": "vi",
                    "raw_segment_count": 1,
                    "cleaned_segment_count": 1,
                    "correction_count": 1,
                    "speaker_assignment_rate": 1.0,
                    "low_confidence_segment_count": 0,
                    "post_processing_applied": True,
                    "quality_status": "healthy",
                }
            },
        )
        db.add(transcript)
        db.commit()
        db.refresh(transcript)

        db.add(
            models.TranscriptSegment(
                transcript_id=transcript.id,
                speaker_label="Speaker_01",
                start_time=0,
                end_time=2,
                text="KPI da duoc chot.",
                original_text="ka pi ai da duoc chot",
                language="vi",
                confidence_score=0.98,
                nlp_metadata={
                    "corrections": [{"wrong": "ka pi ai", "right": "KPI", "source": "rule"}],
                    "speaker_source": "provider",
                    "speaker_confidence": 0.98,
                },
            )
        )
        db.commit()
        meeting_id = meeting.id
    finally:
        db.close()

    response = client.get(f"/api/meetings/{meeting_id}", headers=auth_context["headers"])

    assert response.status_code == 200
    payload = response.json()
    assert payload["transcript_content"] == "KPI da duoc chot."
    assert payload["cleaned_transcript_content"] == "KPI da duoc chot."
    assert payload["raw_transcript_content"] == "ka pi ai da duoc chot"
    assert payload["cleaned_transcript_segments"][0]["text"] == "KPI da duoc chot."
    assert payload["raw_transcript_segments"][0]["text"] == "ka pi ai da duoc chot"
    assert payload["cleaned_transcript_segments"][0]["speaker_source"] == "provider"
    assert payload["cleaned_transcript_segments"][0]["speaker_confidence"] == 0.98
    assert payload["cleaned_transcript_segments"][0]["corrections"][0]["right"] == "KPI"
    assert payload["transcript_quality_metadata"]["quality_status"] == "healthy"


def test_meeting_detail_prefers_current_user_language(client, auth_context):
    db = TestingSessionLocal()
    try:
        english_user = db.query(models.User).filter(models.User.email == "test@example.com").first()
        english_user.language = "en"
        db.commit()

        meeting = models.Meeting(
            organization_id=auth_context["org_id"],
            title="Language Preference",
            description="Meeting with multi-language summaries",
            meeting_type="MEETING",
            status="completed",
            created_by=english_user.id,
        )
        db.add(meeting)
        db.commit()
        db.refresh(meeting)
        meeting_id = meeting.id

        db.add_all(
            [
                models.MeetingSummary(
                    meeting_id=meeting_id,
                    language="vi",
                    meeting_summary="Tom tat tieng Viet",
                    key_points=["vi-point"],
                    decisions=[],
                    processing_status="COMPLETED",
                ),
                models.MeetingSummary(
                    meeting_id=meeting_id,
                    language="en",
                    meeting_summary="English summary",
                    key_points=["en-point"],
                    decisions=[],
                    processing_status="COMPLETED",
                ),
            ]
        )
        db.commit()
    finally:
        db.close()

    response = client.get(f"/api/meetings/{meeting_id}", headers=auth_context["headers"])

    assert response.status_code == 200
    payload = response.json()
    assert payload["preferred_summary_language"] == "en"
    assert payload["meeting_default_summary_language"] == "vi"
    assert payload["meeting_summary_text"] == "English summary"
    assert sorted(payload["available_summary_languages"]) == ["en", "vi"]
    assert payload["summary_generation_state"]["en"] == "COMPLETED"


def test_finalize_defaults_to_user_language_and_does_not_create_ai_tasks(client, auth_context):
    db = TestingSessionLocal()
    try:
        user = db.query(models.User).filter(models.User.email == "test@example.com").first()
        user.language = "en"
        db.commit()
    finally:
        db.close()

    create_response = client.post(
        "/api/meetings",
        headers=auth_context["headers"],
        json=meeting_payload(auth_context, title="Finalize In English", status="processing"),
    )
    meeting_id = create_response.json()["id"]

    router_response = """
    {
      "meeting_summary": "English summary generated from preference.",
      "key_points": ["English point"],
      "decisions": ["English decision"],
      "action_items": [{"task": "Do not persist by default", "owner": "Backend", "deadline": "2026-05-20"}],
      "risks": [],
      "open_questions": [],
      "timeline_highlights": [],
      "speaker_summaries": []
    }
    """

    with patch.dict(
        "os.environ",
        {"ROUTER_API_URL": "http://router.test", "ROUTER_API_KEY": "secret", "ROUTER_MODEL": "router-model"},
        clear=False,
    ), patch(
        "src.providers.router_llm.RouterLLMAdapter.structured_completion",
        return_value=router_response,
    ):
        finalize_response = client.post(
            f"/api/meetings/{meeting_id}/finalize",
            headers=auth_context["headers"],
            json={
                "transcript": "Meeting transcript to summarize in English.",
                "segments": [{"speaker": "Speaker_01", "start": 0, "end": 5, "text": "Hello everyone"}],
                "full_regenerate": False,
            },
        )

    assert finalize_response.status_code == 200
    payload = finalize_response.json()
    assert payload["summary_status"] == "COMPLETED"
    assert payload["default_summary_language"] == "en"
    assert payload["canonical_summary_language"] == "en"
    assert payload["summary"]["meeting_summary"] == "English summary generated from preference."

    db = TestingSessionLocal()
    try:
        summaries = db.query(models.MeetingSummary).filter(models.MeetingSummary.meeting_id == meeting_id).all()
        assert any(summary.language == "en" for summary in summaries)
        action_items = db.query(models.ActionItem).filter(models.ActionItem.meeting_id == meeting_id).all()
        assert action_items == []
    finally:
        db.close()


def test_finalize_full_regenerate_summarizes_once_and_translates_remaining_languages(client, auth_context):
    create_response = client.post(
        "/api/meetings",
        headers=auth_context["headers"],
        json=meeting_payload(auth_context, title="Canonical Summary", status="processing"),
    )
    meeting_id = create_response.json()["id"]

    canonical_response = """
    {
      "meeting_summary": "Canonical English summary.",
      "key_points": ["Point one", "Point two"],
      "decisions": ["Decision one"],
      "action_items": [{"task": "Follow up", "owner": "PM", "deadline": ""}],
      "risks": ["Risk one"],
      "open_questions": ["Question one"],
      "timeline_highlights": ["Kickoff", "Wrap-up"],
      "speaker_summaries": ["PM: aligned the team"]
    }
    """
    translation_response = """
    {
      "meeting_summary": "Ban dich tu ban goc.",
      "key_points": ["Diem mot", "Diem hai"],
      "decisions": ["Quyet dinh mot"],
      "action_items": [{"task": "Theo doi", "owner": "PM", "deadline": ""}],
      "risks": ["Rui ro mot"],
      "open_questions": ["Cau hoi mot"],
      "timeline_highlights": ["Bat dau", "Ket thuc"],
      "speaker_summaries": ["PM: da dong bo doi ngu"]
    }
    """
    prompts = []

    def fake_structured_completion(self, system_prompt, user_prompt, temperature=0.1, max_tokens=2000):
        prompts.append(user_prompt)
        if "Canonical summary JSON:" in user_prompt:
            return translation_response
        return canonical_response

    with patch.dict(
        "os.environ",
        {"ROUTER_API_URL": "http://router.test", "ROUTER_API_KEY": "secret", "ROUTER_MODEL": "router-model"},
        clear=False,
    ), patch(
        "src.providers.router_llm.RouterLLMAdapter.structured_completion",
        new=fake_structured_completion,
    ):
        finalize_response = client.post(
            f"/api/meetings/{meeting_id}/finalize",
            headers=auth_context["headers"],
            json={
                "transcript": "Transcript to summarize once and translate later.",
                "segments": [{"speaker": "Speaker_01", "start": 0, "end": 5, "text": "Hello everyone"}],
                "language": "en",
                "regenerate": True,
                "full_regenerate": True,
                "generate_action_items": False,
            },
        )

    assert finalize_response.status_code == 200
    payload = finalize_response.json()
    assert payload["summary_status"] == "COMPLETED"
    assert payload["translation_queue_started"] is True
    assert payload["canonical_summary_language"] == "en"
    assert payload["summary_generation_state"]["en"] == "COMPLETED"
    assert sum("Transcript:" in prompt for prompt in prompts) == 1
    assert sum("Canonical summary JSON:" in prompt for prompt in prompts) == 4

    db = TestingSessionLocal()
    try:
        summaries = db.query(models.MeetingSummary).filter(models.MeetingSummary.meeting_id == meeting_id).all()
        assert sorted(summary.language for summary in summaries if summary.processing_status == "COMPLETED") == ["en", "ja", "ko", "vi", "zh"]
        counts = {(summary.language, len(summary.key_points or []), len(summary.decisions or [])) for summary in summaries}
        assert ("en", 2, 1) in counts
        assert ("vi", 2, 1) in counts
    finally:
        db.close()


def test_get_meeting_includes_canonical_summary_id_and_keeps_ai_tasks_for_translations(client, auth_context):
    db = TestingSessionLocal()
    try:
        user = db.query(models.User).filter(models.User.email == "test@example.com").first()
        user.language = "vi"
        meeting = models.Meeting(
            organization_id=auth_context["org_id"],
            title="Canonical Tasks",
            description="Translation should still see canonical AI tasks",
            meeting_type="MEETING",
            status="completed",
            created_by=user.id,
        )
        db.add(meeting)
        db.commit()
        db.refresh(meeting)

        canonical_summary = models.MeetingSummary(
            meeting_id=meeting.id,
            language="en",
            meeting_summary="Canonical summary",
            key_points=["canonical"],
            decisions=[],
            action_items=[{"task": "Canonical AI task", "owner": "PM", "deadline": ""}],
            processing_status="COMPLETED",
            generation_group_id="group-1",
            summary_kind="canonical",
        )
        translation_summary = models.MeetingSummary(
            meeting_id=meeting.id,
            language="vi",
            meeting_summary="Ban dich",
            key_points=["ban dich"],
            decisions=[],
            action_items=[{"task": "Tac vu AI", "owner": "PM", "deadline": ""}],
            processing_status="COMPLETED",
            generation_group_id="group-1",
            summary_kind="translation",
        )
        db.add_all([canonical_summary, translation_summary])
        db.commit()
        db.refresh(canonical_summary)

        db.add(
            models.ActionItem(
                meeting_id=meeting.id,
                summary_id=canonical_summary.id,
                title="Canonical AI task",
                status="PENDING",
                priority="MEDIUM",
                created_by=user.id,
            )
        )
        db.commit()
        canonical_summary_id = canonical_summary.id
        meeting_id = meeting.id
    finally:
        db.close()

    response = client.get(f"/api/meetings/{meeting_id}", headers=auth_context["headers"])

    assert response.status_code == 200
    payload = response.json()
    assert payload["canonical_summary_language"] == "en"
    assert payload["canonical_summary_id"] == canonical_summary_id
    assert payload["generation_group_id"] == "group-1"
    assert len(payload["action_items"]) == 1
    assert payload["action_items"][0]["summary_id"] == canonical_summary_id


def test_get_meetings_uses_latest_generation_completed_summary_preview_only(client, auth_context):
    db = TestingSessionLocal()
    try:
        user = db.query(models.User).filter(models.User.email == "test@example.com").first()
        user.language = "vi"
        meeting = models.Meeting(
            organization_id=auth_context["org_id"],
            title="Preview Meeting",
            description="Meeting list preview",
            meeting_type="MEETING",
            status="completed",
            created_by=user.id,
        )
        db.add(meeting)
        db.commit()
        db.refresh(meeting)

        db.add_all(
            [
                models.MeetingSummary(
                    meeting_id=meeting.id,
                    language="vi",
                    meeting_summary="Dang dich tu ban goc.",
                    key_points=[],
                    decisions=[],
                    processing_status="PROCESSING",
                    generation_group_id="group-new",
                    summary_kind="translation",
                ),
                models.MeetingSummary(
                    meeting_id=meeting.id,
                    language="en",
                    meeting_summary="Canonical latest summary",
                    key_points=["latest point"],
                    decisions=["latest decision"],
                    processing_status="COMPLETED",
                    generation_group_id="group-new",
                    summary_kind="canonical",
                ),
                models.MeetingSummary(
                    meeting_id=meeting.id,
                    language="vi",
                    meeting_summary="Old stale summary",
                    key_points=["old point"],
                    decisions=["old decision"],
                    processing_status="COMPLETED",
                    generation_group_id="group-old",
                    summary_kind="canonical",
                ),
            ]
        )
        db.commit()
    finally:
        db.close()

    response = client.get("/api/meetings", headers=auth_context["headers"])

    assert response.status_code == 200
    payload = response.json()
    preview = next(item for item in payload if item["title"] == "Preview Meeting")
    assert preview["summary_text"] == "Canonical latest summary"
    assert preview["key_points_list"] == ["latest point"]
    assert preview["decisions_list"] == ["latest decision"]


def test_finalize_broadcasts_translation_events_and_returns_rate_limit_error(client, auth_context):
    create_response = client.post(
        "/api/meetings",
        headers=auth_context["headers"],
        json=meeting_payload(auth_context, title="Broadcast Summary", status="processing"),
    )
    meeting_id = create_response.json()["id"]

    canonical_response = """
    {
      "meeting_summary": "Canonical English summary.",
      "key_points": ["Point one"],
      "decisions": ["Decision one"],
      "action_items": [],
      "risks": [],
      "open_questions": [],
      "timeline_highlights": [],
      "speaker_summaries": []
    }
    """
    translation_response = """
    {
      "meeting_summary": "Ban dich tu ban goc.",
      "key_points": ["Diem mot"],
      "decisions": ["Quyet dinh mot"],
      "action_items": [],
      "risks": [],
      "open_questions": [],
      "timeline_highlights": [],
      "speaker_summaries": []
    }
    """
    events = []

    async def capture_broadcast(meeting_id, payload):
        events.append(payload)

    def fake_structured_completion(self, system_prompt, user_prompt, temperature=0.1, max_tokens=2000):
        if "Canonical summary JSON:" in user_prompt:
            if "Vietnamese" in system_prompt:
                self.last_error = "Groq rate limit exceeded"
                self.last_error_type = "rate_limit"
                self.last_retry_after_seconds = 0.01
                return None
            return translation_response
        return canonical_response

    with patch.dict(
        "os.environ",
        {"ROUTER_API_URL": "http://router.test", "ROUTER_API_KEY": "secret", "ROUTER_MODEL": "router-model"},
        clear=False,
    ), patch(
        "src.providers.router_llm.RouterLLMAdapter.structured_completion",
        new=fake_structured_completion,
    ), patch(
        "src.api.core.meeting_operations.meeting_room_manager.broadcast",
        new=capture_broadcast,
    ):
        finalize_response = client.post(
            f"/api/meetings/{meeting_id}/finalize",
            headers=auth_context["headers"],
            json={
                "transcript": "Transcript to summarize once and translate later.",
                "segments": [{"speaker": "Speaker_01", "start": 0, "end": 5, "text": "Hello everyone"}],
                "language": "en",
                "regenerate": True,
                "full_regenerate": True,
                "generate_action_items": False,
            },
        )

    assert finalize_response.status_code == 200
    completed_languages = {
        event["language"]
        for event in events
        if event["type"] == "ai.notes.completed" and event["summary_status"] == "COMPLETED"
    }
    failed_languages = {
        event["language"]
        for event in events
        if event["type"] == "ai.notes.failed"
    }
    started_languages = {
        event["language"]
        for event in events
        if event["type"] == "ai.notes.started"
    }

    assert "en" in completed_languages
    assert "ja" in started_languages or "ko" in started_languages or "zh" in started_languages or "vi" in started_languages
    assert "vi" in failed_languages
    failed_vi = next(event for event in events if event["type"] == "ai.notes.failed" and event["language"] == "vi")
    assert failed_vi["generation_group_id"]
    assert failed_vi["error_type"] == "rate_limit"
    assert "rate limit" in failed_vi["error"].lower()


def test_regenerate_with_ai_tasks_replaces_only_ai_generated_tasks(client, auth_context):
    db = TestingSessionLocal()
    try:
        user = db.query(models.User).filter(models.User.email == "test@example.com").first()
        user.language = "en"
        meeting = models.Meeting(
            organization_id=auth_context["org_id"],
            title="Regenerate Tasks",
            description="Task replacement behavior",
            meeting_type="MEETING",
            status="completed",
            created_by=user.id,
        )
        db.add(meeting)
        db.commit()
        db.refresh(meeting)
        meeting_id = meeting.id

        summary = models.MeetingSummary(
            meeting_id=meeting_id,
            language="en",
            meeting_summary="Old summary",
            key_points=["old"],
            decisions=[],
            action_items=[{"task": "Old AI task", "owner": "Backend", "deadline": ""}],
            processing_status="COMPLETED",
        )
        db.add(summary)
        db.commit()
        db.refresh(summary)

        db.add(
            models.ActionItem(
                meeting_id=meeting_id,
                summary_id=summary.id,
                title="Old AI task",
                status="PENDING",
                priority="MEDIUM",
                created_by=user.id,
            )
        )
        db.add(
            models.ActionItem(
                meeting_id=meeting_id,
                summary_id=None,
                title="Manual task",
                status="PENDING",
                priority="HIGH",
                created_by=user.id,
            )
        )
        db.commit()
    finally:
        db.close()

    router_response = """
    {
      "meeting_summary": "New summary",
      "key_points": ["new"],
      "decisions": ["new decision"],
      "action_items": [{"task": "New AI task", "owner": "Backend", "deadline": "2026-05-20"}],
      "risks": [],
      "open_questions": [],
      "timeline_highlights": [],
      "speaker_summaries": []
    }
    """

    with patch.dict(
        "os.environ",
        {"ROUTER_API_URL": "http://router.test", "ROUTER_API_KEY": "secret", "ROUTER_MODEL": "router-model"},
        clear=False,
    ), patch(
        "src.providers.router_llm.RouterLLMAdapter.structured_completion",
        return_value=router_response,
    ):
        response = client.post(
            f"/api/meetings/{meeting_id}/finalize",
            headers=auth_context["headers"],
            json={
                "transcript": "Transcript for new summary",
                "segments": [{"speaker": "Speaker_01", "start": 0, "end": 5, "text": "Need a new task"}],
                "language": "en",
                "regenerate": True,
                "full_regenerate": False,
                "generate_action_items": True,
            },
        )

    assert response.status_code == 200
    db = TestingSessionLocal()
    try:
        items = db.query(models.ActionItem).filter(models.ActionItem.meeting_id == meeting_id).all()
        titles = sorted(item.title for item in items)
        assert titles == ["Manual task", "New AI task"]
    finally:
        db.close()


def test_regenerate_without_ai_tasks_preserves_all_existing_tasks(client, auth_context):
    db = TestingSessionLocal()
    try:
        user = db.query(models.User).filter(models.User.email == "test@example.com").first()
        user.language = "en"
        meeting = models.Meeting(
            organization_id=auth_context["org_id"],
            title="Preserve Tasks",
            description="Regenerate without AI tasks should keep everything",
            meeting_type="MEETING",
            status="completed",
            created_by=user.id,
        )
        db.add(meeting)
        db.commit()
        db.refresh(meeting)
        meeting_id = meeting.id

        summary = models.MeetingSummary(
            meeting_id=meeting_id,
            language="en",
            meeting_summary="Old summary",
            key_points=["old"],
            decisions=[],
            action_items=[{"task": "Old AI task", "owner": "Backend", "deadline": ""}],
            processing_status="COMPLETED",
        )
        db.add(summary)
        db.commit()
        db.refresh(summary)

        db.add(
            models.ActionItem(
                meeting_id=meeting_id,
                summary_id=summary.id,
                title="Old AI task",
                status="PENDING",
                priority="MEDIUM",
                created_by=user.id,
            )
        )
        db.add(
            models.ActionItem(
                meeting_id=meeting_id,
                summary_id=None,
                title="Manual task",
                status="PENDING",
                priority="HIGH",
                created_by=user.id,
            )
        )
        db.commit()
    finally:
        db.close()

    router_response = """
    {
      "meeting_summary": "New summary without action items",
      "key_points": ["new"],
      "decisions": ["new decision"],
      "action_items": [{"task": "Should not be persisted", "owner": "Backend", "deadline": "2026-05-20"}],
      "risks": [],
      "open_questions": [],
      "timeline_highlights": [],
      "speaker_summaries": []
    }
    """

    with patch.dict(
        "os.environ",
        {"ROUTER_API_URL": "http://router.test", "ROUTER_API_KEY": "secret", "ROUTER_MODEL": "router-model"},
        clear=False,
    ), patch(
        "src.providers.router_llm.RouterLLMAdapter.structured_completion",
        return_value=router_response,
    ):
        response = client.post(
            f"/api/meetings/{meeting_id}/finalize",
            headers=auth_context["headers"],
            json={
                "transcript": "Transcript for regenerate without AI tasks",
                "segments": [{"speaker": "Speaker_01", "start": 0, "end": 5, "text": "Keep all tasks"}],
                "language": "en",
                "regenerate": True,
                "full_regenerate": False,
                "generate_action_items": False,
            },
        )

    assert response.status_code == 200
    payload = response.json()
    assert payload["summary_status"] == "COMPLETED"
    assert payload["default_summary_language"] == "en"

    db = TestingSessionLocal()
    try:
        items = db.query(models.ActionItem).filter(models.ActionItem.meeting_id == meeting_id).all()
        titles = sorted(item.title for item in items)
        assert titles == ["Manual task", "Old AI task"]
        summary_db = (
            db.query(models.MeetingSummary)
            .filter(models.MeetingSummary.meeting_id == meeting_id, models.MeetingSummary.language == "en")
            .order_by(models.MeetingSummary.created_at.desc())
            .first()
        )
        assert summary_db is not None
        assert summary_db.meeting_summary == "New summary without action items"
        assert summary_db.action_items == [{"task": "Should not be persisted", "owner": "Backend", "deadline": "2026-05-20"}]
    finally:
        db.close()


def test_export_meeting_docx_minutes_with_transcript_appendix(client, auth_context):
    create_response = client.post(
        "/api/meetings",
        headers=auth_context["headers"],
        json=meeting_payload(auth_context, title="Export Minutes"),
    )
    meeting_id = create_response.json()["id"]

    db = TestingSessionLocal()
    try:
        user = db.query(models.User).filter(models.User.email == "test@example.com").first()
        db.add(
            models.MeetingParticipant(
                meeting_id=meeting_id,
                email="alice@example.com",
                name="Alice Nguyen",
                role="PARTICIPANT",
            )
        )
        transcript = models.Transcript(
            meeting_id=meeting_id,
            content="Chung ta thong nhat rollout he thong moi.",
            language="vi",
            processing_status="COMPLETED",
            stt_provider="deepgram",
        )
        db.add(transcript)
        db.flush()
        db.add(
            models.TranscriptSegment(
                transcript_id=transcript.id,
                speaker_label="Speaker_01",
                start_time=0,
                end_time=5,
                text="Chung ta thong nhat rollout he thong moi.",
                language="vi",
            )
        )
        summary = models.MeetingSummary(
            meeting_id=meeting_id,
            language="vi",
            meeting_summary="Cuoc hop thong nhat rollout he thong moi.",
            key_points=["Thong nhat timeline", "Chot owner"],
            decisions=["Go live vao thu Hai"],
            action_items=["Alice Nguyen cap nhat checklist"],
            processing_status="COMPLETED",
        )
        db.add(summary)
        db.flush()
        db.add(
            models.ActionItem(
                meeting_id=meeting_id,
                summary_id=summary.id,
                title="Cap nhat checklist rollout",
                created_by=user.id,
                assigned_email="alice@example.com",
                status="PENDING",
                priority="HIGH",
            )
        )
        db.commit()
    finally:
        db.close()

    response = client.post(
        "/api/export/generate",
        headers=auth_context["headers"],
        json={
            "meeting_id": meeting_id,
            "format": "docx",
            "language": "vi",
            "include_transcript_appendix": True,
            "include_summary": True,
            "include_action_items": True,
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["filename"].startswith("export-minutes_")
    assert payload["filename"].endswith(".docx")

    download_response = client.get(payload["download_url"], headers=auth_context["headers"])
    assert download_response.status_code == 200

    document = Document(BytesIO(download_response.content))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "BIÊN BẢN CUỘC HỌP" in text
    assert "Export Minutes" in text
    assert "Tóm tắt cuộc họp" in text
    assert "PHỤ LỤC A. TRANSCRIPT CUỘC HỌP" in text
    assert "Chung ta thong nhat rollout he thong moi." in text


def test_export_meeting_pdf_without_summary_still_succeeds(client, auth_context):
    create_response = client.post(
        "/api/meetings",
        headers=auth_context["headers"],
        json=meeting_payload(auth_context, title="PDF Export", status="completed"),
    )
    meeting_id = create_response.json()["id"]

    response = client.post(
        "/api/export/generate",
        headers=auth_context["headers"],
        json={
            "meeting_id": meeting_id,
            "format": "pdf",
            "language": "en",
            "include_transcript_appendix": False,
            "include_summary": True,
            "include_action_items": True,
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["filename"].endswith(".pdf")

    download_response = client.get(payload["download_url"], headers=auth_context["headers"])
    assert download_response.status_code == 200
    assert download_response.headers["content-type"] == "application/pdf"
    assert download_response.content.startswith(b"%PDF")


def test_export_download_rejects_invalid_filename_format(client, auth_context):
    response = client.get(
        "/api/export/download/bien-ban-cuoc-hop-not-a-real-id-20260523_161753.pdf",
        headers=auth_context["headers"],
    )

    assert response.status_code == 400
    assert response.json()["detail"] == "Invalid filename"


def test_export_placeholder_title_uses_clean_fallback_and_creator_participant(client, auth_context):
    create_response = client.post(
        "/api/meetings",
        headers=auth_context["headers"],
        json=meeting_payload(
            auth_context,
            title="cuộc họp",
            scheduled_start="2026-05-27T09:00:00Z",
        ),
    )
    meeting_id = create_response.json()["id"]

    response = client.post(
        "/api/export/generate",
        headers=auth_context["headers"],
        json={
            "meeting_id": meeting_id,
            "format": "docx",
            "language": "vi",
            "include_transcript_appendix": False,
            "include_summary": True,
            "include_action_items": True,
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["filename"].startswith("cuoc-hop_")
    assert payload["filename"].endswith(".docx")

    download_response = client.get(payload["download_url"], headers=auth_context["headers"])
    assert download_response.status_code == 200

    document = Document(BytesIO(download_response.content))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Tên cuộc họp: cuộc họp" in text
    assert "testuser - test@example.com (Người tạo cuộc họp)" in text
    assert "HOST" not in text
    assert "N/A" not in text
