import pytest
from unittest.mock import patch
from io import BytesIO
from pathlib import Path
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from sqlalchemy.pool import StaticPool
from docx import Document
from src.api.main import app
from src.api.database import get_db, Base
from src.api import models

TEST_DATABASE_URL = "sqlite:///:memory:"
test_engine = create_engine(
    TEST_DATABASE_URL,
    connect_args={"check_same_thread": False},
    poolclass=StaticPool,
)
TestingSessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=test_engine)


def override_get_db():
    db = TestingSessionLocal()
    try:
        yield db
    finally:
        db.close()


app.dependency_overrides[get_db] = override_get_db


def activate_org(org_id: str):
    db = TestingSessionLocal()
    try:
        org = db.query(models.Organization).filter(models.Organization.id == org_id).first()
        org.settings = {"approval_status": "active"}
        db.commit()
    finally:
        db.close()


@pytest.fixture(scope="function")
def client():
    Base.metadata.create_all(bind=test_engine)
    app.dependency_overrides[get_db] = override_get_db
    with TestClient(app) as test_client:
        yield test_client
    app.dependency_overrides.pop(get_db, None)
    Base.metadata.drop_all(bind=test_engine)


@pytest.fixture
def auth_context(client):
    client.post(
        "/api/auth/register",
        json={
            "username": "testuser",
            "email": "test@example.com",
            "password": "securepassword123",
            "orgName": "Test Organization",
        },
    )

    response = client.post(
        "/api/auth/login",
        json={
            "username": "testuser",
            "password": "securepassword123",
        },
    )
    data = response.json()
    activate_org(data["user"]["orgMemberships"][0]["orgId"])
    return {
        "headers": {"Authorization": f"Bearer {data['access_token']}"},
        "org_id": data["user"]["orgMemberships"][0]["orgId"],
    }


def meeting_payload(auth_context, **overrides):
    data = {
        "organization_id": auth_context["org_id"],
        "title": "Test Meeting",
        "description": "A test meeting",
        "meeting_type": "MEETING",
        "status": "upcoming",
    }
    data.update(overrides)
    return data


def test_create_meeting(client, auth_context):
    response = client.post(
        "/api/meetings",
        headers=auth_context["headers"],
        json=meeting_payload(auth_context),
    )

    assert response.status_code == 200
    data = response.json()
    assert data["title"] == "Test Meeting"
    assert data["status"] == "upcoming"
    assert data["meeting_type"] == "MEETING"
    assert data["organization_id"] == auth_context["org_id"]
    assert "id" in data


def test_create_meeting_unauthorized(client):
    response = client.post(
        "/api/meetings",
        json={"title": "Test Meeting"},
    )

    assert response.status_code == 401


def test_get_meetings(client, auth_context):
    client.post(
        "/api/meetings",
        headers=auth_context["headers"],
        json=meeting_payload(auth_context, title="Test Meeting 1"),
    )
    client.post(
        "/api/meetings",
        headers=auth_context["headers"],
        json=meeting_payload(auth_context, title="Test Meeting 2"),
    )

    response = client.get("/api/meetings", headers=auth_context["headers"])

    assert response.status_code == 200
    data = response.json()
    assert isinstance(data, list)
    assert len(data) >= 2


def test_get_meeting_by_id(client, auth_context):
    create_response = client.post(
        "/api/meetings",
        headers=auth_context["headers"],
        json=meeting_payload(auth_context),
    )
    meeting_id = create_response.json()["id"]

    response = client.get(f"/api/meetings/{meeting_id}", headers=auth_context["headers"])

    assert response.status_code == 200
    data = response.json()
    assert data["id"] == meeting_id
    assert data["title"] == "Test Meeting"


def test_get_meeting_not_found(client, auth_context):
    response = client.get("/api/meetings/nonexistent-id", headers=auth_context["headers"])

    assert response.status_code == 404
    assert "not found" in response.json()["detail"].lower()


def test_update_meeting(client, auth_context):
    create_response = client.post(
        "/api/meetings",
        headers=auth_context["headers"],
        json=meeting_payload(auth_context, title="Original Title"),
    )
    meeting_id = create_response.json()["id"]

    response = client.put(
        f"/api/meetings/{meeting_id}",
        headers=auth_context["headers"],
        json={
            "title": "Updated Title",
            "status": "processing",
            "description": "Updated description",
        },
    )

    assert response.status_code == 200
    data = response.json()
    assert data["title"] == "Updated Title"
    assert data["status"] == "processing"
    assert data["description"] == "Updated description"


def test_delete_meeting(client, auth_context):
    create_response = client.post(
        "/api/meetings",
        headers=auth_context["headers"],
        json=meeting_payload(auth_context),
    )
    meeting_id = create_response.json()["id"]

    response = client.delete(f"/api/meetings/{meeting_id}", headers=auth_context["headers"])

    assert response.status_code == 200
    assert "deleted successfully" in response.json()["message"].lower()

    get_response = client.get(f"/api/meetings/{meeting_id}", headers=auth_context["headers"])
    assert get_response.status_code == 404


def test_get_meetings_with_filters(client, auth_context):
    client.post(
        "/api/meetings",
        headers=auth_context["headers"],
        json=meeting_payload(auth_context, title="Upcoming Meeting", status="upcoming"),
    )
    client.post(
        "/api/meetings",
        headers=auth_context["headers"],
        json=meeting_payload(auth_context, title="Completed Meeting", status="completed"),
    )

    response = client.get("/api/meetings?status=upcoming", headers=auth_context["headers"])

    assert response.status_code == 200
    data = response.json()
    assert all(m["status"] == "upcoming" for m in data)


def test_upload_audio_creates_meeting_audio_and_job(client, auth_context, monkeypatch, tmp_path):
    from src.api.routes import stt as stt_routes

    monkeypatch.setattr(stt_routes, "ensure_audio_upload_dir", lambda: str(tmp_path))
    monkeypatch.setattr(stt_routes, "start_upload_job", lambda job: None)

    response = client.post(
        "/api/upload",
        headers=auth_context["headers"],
        files={"file": ("backup.wav", b"fake-wave-data", "audio/wav")},
        data={
            "organization_id": auth_context["org_id"],
            "title": "Backup Upload",
            "language": "auto",
            "stt_provider": "deepgram",
            "enable_diarization": "true",
            "enable_glossary": "true",
            "enable_summary": "true",
            "enable_action_items": "true",
            "enable_noise_cleanup": "true",
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["status"] == "queued"
    assert payload["current_stage"] == "queued"
    assert payload["progress_percent"] == 0

    status_response = client.get(f"/api/upload/jobs/{payload['job_id']}", headers=auth_context["headers"])
    assert status_response.status_code == 200
    assert status_response.json()["meeting_id"] == payload["meeting_id"]

    db = TestingSessionLocal()
    try:
        meeting = db.query(models.Meeting).filter(models.Meeting.id == payload["meeting_id"]).first()
        assert meeting is not None
        assert meeting.status == "queued"
        assert meeting.audio_url
        assert meeting.recording_url

        audio_file = db.query(models.AudioFile).filter(models.AudioFile.meeting_id == meeting.id).first()
        assert audio_file is not None
        assert audio_file.upload_status == "UPLOADED"
        assert Path(audio_file.file_path).exists()
    finally:
        db.close()


def test_get_meetings_unauthorized(client):
    response = client.get("/api/meetings")

    assert response.status_code == 401


def test_finalize_meeting_single_router_call_success(client, auth_context):
    create_response = client.post(
        "/api/meetings",
        headers=auth_context["headers"],
        json=meeting_payload(auth_context, title="Finalize Success", status="processing"),
    )
    meeting_id = create_response.json()["id"]

    router_response = """
    {
      "meeting_summary": "Cuoc hop thong nhat chot scope STT va summary.",
      "key_points": ["Chot luong live chunk", "Dung Router cho summary"],
      "decisions": ["Bo 4 call rieng trong finalize"],
      "action_items": [{"task": "Cap nhat finalize", "owner": "Backend", "deadline": "2026-05-20"}]
    }
    """

    with patch.dict(
        "os.environ",
        {"ROUTER_API_URL": "http://router.test", "ROUTER_API_KEY": "secret", "ROUTER_MODEL": "router-model"},
        clear=False,
    ), patch(
        "src.providers.router_llm.RouterLLMAdapter.structured_completion",
        return_value=router_response,
    ) as structured_completion:
        finalize_response = client.post(
            f"/api/meetings/{meeting_id}/finalize",
            headers=auth_context["headers"],
            json={
                "transcript": "Hom nay chung ta chot luong live chunk va dung Router de tom tat.",
                "segments": [{"speaker": "Speaker_01", "start": 0, "end": 5, "text": "Noi dung hop"}],
            },
        )

    assert finalize_response.status_code == 200
    data = finalize_response.json()
    assert data["meeting_id"] == meeting_id
    assert data["transcript_status"] == "COMPLETED"
    assert data["summary_status"] == "COMPLETED"
    assert data["summary"]["meeting_summary"] == "Cuoc hop thong nhat chot scope STT va summary."
    assert data["summary"]["key_points"] == ["Chot luong live chunk", "Dung Router cho summary"]
    assert len(data["summary"]["action_items"]) == 1
    assert data["errors"] == []
    structured_completion.assert_called_once()

    detail_response = client.get(f"/api/meetings/{meeting_id}", headers=auth_context["headers"])
    assert detail_response.status_code == 200
    detail = detail_response.json()
    assert detail["transcript_content"] == "Hom nay chung ta chot luong live chunk va dung Router de tom tat."
    assert detail["meeting_summary_text"] == "Cuoc hop thong nhat chot scope STT va summary."
    assert detail["key_points_text"] == ["Chot luong live chunk", "Dung Router cho summary"]
    assert detail["decisions_text"] == ["Bo 4 call rieng trong finalize"]
    assert len(detail["action_items"]) == 1
    assert detail["action_items"][0]["title"] == "Cap nhat finalize"


def test_finalize_meeting_router_failure_marks_summary_failed(client, auth_context):
    create_response = client.post(
        "/api/meetings",
        headers=auth_context["headers"],
        json=meeting_payload(auth_context, title="Finalize Failure", status="processing"),
    )
    meeting_id = create_response.json()["id"]

    with patch.dict(
        "os.environ",
        {"ROUTER_API_URL": "http://router.test", "ROUTER_API_KEY": "secret", "ROUTER_MODEL": "router-model"},
        clear=False,
    ), patch(
        "src.providers.router_llm.RouterLLMAdapter.structured_completion",
        return_value="not valid json",
    ):
        finalize_response = client.post(
            f"/api/meetings/{meeting_id}/finalize",
            headers=auth_context["headers"],
            json={
                "transcript": "Transcript van duoc luu du Router tra ve sai dinh dang.",
                "segments": [],
            },
        )

    assert finalize_response.status_code == 200
    data = finalize_response.json()
    assert data["transcript_status"] == "COMPLETED"
    assert data["summary_status"] == "FAILED"
    assert data["summary"]["meeting_summary"] == ""
    assert data["summary"]["key_points"] == []
    assert data["summary"]["decisions"] == []
    assert data["errors"]

    detail_response = client.get(f"/api/meetings/{meeting_id}", headers=auth_context["headers"])
    assert detail_response.status_code == 200
    detail = detail_response.json()
    assert detail["transcript_content"] == "Transcript van duoc luu du Router tra ve sai dinh dang."
    assert detail["meeting_summary_text"] is None
    assert detail["key_points_text"] == []
    assert detail["decisions_text"] == []
    assert detail["action_items"] == []


def test_export_meeting_docx_minutes_with_transcript_appendix(client, auth_context):
    create_response = client.post(
        "/api/meetings",
        headers=auth_context["headers"],
        json=meeting_payload(auth_context, title="Export Minutes", status="completed"),
    )
    meeting_id = create_response.json()["id"]

    db = TestingSessionLocal()
    try:
        user = db.query(models.User).filter(models.User.email == "test@example.com").first()
        db.add(
            models.MeetingParticipant(
                meeting_id=meeting_id,
                email="alice@example.com",
                name="Alice Nguyen",
                role="PARTICIPANT",
            )
        )
        transcript = models.Transcript(
            meeting_id=meeting_id,
            content="Chung ta thong nhat rollout he thong moi.",
            language="vi",
            processing_status="COMPLETED",
            stt_provider="deepgram",
        )
        db.add(transcript)
        db.flush()
        db.add(
            models.TranscriptSegment(
                transcript_id=transcript.id,
                speaker_label="Speaker_01",
                start_time=0,
                end_time=5,
                text="Chung ta thong nhat rollout he thong moi.",
                language="vi",
            )
        )
        summary = models.MeetingSummary(
            meeting_id=meeting_id,
            language="vi",
            meeting_summary="Cuoc hop thong nhat rollout he thong moi.",
            key_points=["Thong nhat timeline", "Chot owner"],
            decisions=["Go live vao thu Hai"],
            action_items=["Alice Nguyen cap nhat checklist"],
            processing_status="COMPLETED",
        )
        db.add(summary)
        db.flush()
        db.add(
            models.ActionItem(
                meeting_id=meeting_id,
                summary_id=summary.id,
                title="Cap nhat checklist rollout",
                created_by=user.id,
                assigned_email="alice@example.com",
                status="PENDING",
                priority="HIGH",
            )
        )
        db.commit()
    finally:
        db.close()

    response = client.post(
        "/api/export/generate",
        headers=auth_context["headers"],
        json={
            "meeting_id": meeting_id,
            "format": "docx",
            "language": "vi",
            "include_transcript_appendix": True,
            "include_summary": True,
            "include_action_items": True,
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["filename"].startswith("bien-ban-export-minutes-")
    assert meeting_id in payload["filename"]
    assert payload["filename"].endswith(".docx")

    download_response = client.get(payload["download_url"], headers=auth_context["headers"])
    assert download_response.status_code == 200

    document = Document(BytesIO(download_response.content))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "BIÊN BẢN CUỘC HỌP" in text
    assert "Export Minutes" in text
    assert "Tóm tắt cuộc họp" in text
    assert "PHỤ LỤC A. TRANSCRIPT CUỘC HỌP" in text
    assert "Chung ta thong nhat rollout he thong moi." in text


def test_export_meeting_pdf_without_summary_still_succeeds(client, auth_context):
    create_response = client.post(
        "/api/meetings",
        headers=auth_context["headers"],
        json=meeting_payload(auth_context, title="PDF Export", status="completed"),
    )
    meeting_id = create_response.json()["id"]

    response = client.post(
        "/api/export/generate",
        headers=auth_context["headers"],
        json={
            "meeting_id": meeting_id,
            "format": "pdf",
            "language": "en",
            "include_transcript_appendix": False,
            "include_summary": True,
            "include_action_items": True,
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["filename"].endswith(".pdf")

    download_response = client.get(payload["download_url"], headers=auth_context["headers"])
    assert download_response.status_code == 200
    assert download_response.headers["content-type"] == "application/pdf"
    assert download_response.content.startswith(b"%PDF")


def test_export_download_rejects_invalid_filename_format(client, auth_context):
    response = client.get(
        "/api/export/download/bien-ban-cuoc-hop-not-a-real-id-20260523_161753.pdf",
        headers=auth_context["headers"],
    )

    assert response.status_code == 400
    assert response.json()["detail"] == "Invalid filename"


def test_export_placeholder_title_uses_clean_fallback_and_creator_participant(client, auth_context):
    create_response = client.post(
        "/api/meetings",
        headers=auth_context["headers"],
        json=meeting_payload(
            auth_context,
            title="cuộc họp",
            status="completed",
            scheduled_start="2026-05-23T09:00:00Z",
        ),
    )
    meeting_id = create_response.json()["id"]

    response = client.post(
        "/api/export/generate",
        headers=auth_context["headers"],
        json={
            "meeting_id": meeting_id,
            "format": "docx",
            "language": "vi",
            "include_transcript_appendix": False,
            "include_summary": True,
            "include_action_items": True,
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["filename"].startswith("bien-ban-chua-cap-nhat-ten-")
    assert meeting_id in payload["filename"]

    download_response = client.get(payload["download_url"], headers=auth_context["headers"])
    assert download_response.status_code == 200

    document = Document(BytesIO(download_response.content))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Chưa cập nhật tên cuộc họp" in text
    assert "Bắt đầu: 23/05/2026 09:00" in text
    assert "Kết thúc: Chưa có dữ liệu" in text
    assert "testuser - test@example.com (Người tạo cuộc họp)" in text
    assert "HOST" not in text
    assert "N/A" not in text
